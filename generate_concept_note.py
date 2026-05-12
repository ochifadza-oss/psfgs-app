"""
PSFGS — SALGA Concept Note Word Document Generator
Glance Management Technologies (Pty) Ltd · 2026
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "SALGA_ConceptNote_PSFGS_July2026.docx")

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x38, 0x64)   # #1F3864
MID_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6
LIGHT_BLUE  = RGBColor(0xEE, 0xF5, 0xFC)   # #EEF5FC  (cell shade)
DARK_GREEN  = RGBColor(0x1A, 0x6B, 0x3C)   # #1A6B3C
AMBER       = RGBColor(0xD4, 0x86, 0x1A)   # #D4861A
RED         = RGBColor(0xC0, 0x39, 0x2B)   # #C0392B
GREY_TEXT   = RGBColor(0x5A, 0x64, 0x78)   # #5A6478
LIGHT_GREY  = RGBColor(0xF4, 0xF6, 0xF9)   # #F4F6F9
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
BODY_TEXT   = RGBColor(0x2C, 0x3E, 0x50)   # #2C3E50


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def hex_colour(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def shade_cell(cell, fill_rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour(fill_rgb))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """kwargs: top/bottom/left/right = (colour_hex, size_eighths, val)"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, (colour, sz, val) in kwargs.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), colour)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top=('FFFFFF', 0, 'none'), bottom=('FFFFFF', 0, 'none'),
                left=('FFFFFF', 0, 'none'), right=('FFFFFF', 0, 'none'))

def thick_left_border(cell, colour_rgb: RGBColor):
    set_cell_border(cell,
        left=(hex_colour(colour_rgb), 24, 'single'),
        top=('FFFFFF', 0, 'none'), bottom=('FFFFFF', 0, 'none'),
        right=('FFFFFF', 0, 'none'))

def cell_spacing(cell, top=40, bottom=40, left=80, right=80):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)

def merge_row(table, row_idx):
    """Merge all cells in a row."""
    row   = table.rows[row_idx]
    cells = row.cells
    cells[0].merge(cells[-1])
    return cells[0]

def paragraph_spacing(para, before=0, after=0, line=None):
    pPr  = para._p.get_or_add_pPr()
    spng = OxmlElement('w:spacing')
    spng.set(qn('w:before'), str(before))
    spng.set(qn('w:after'),  str(after))
    if line:
        spng.set(qn('w:line'),     str(line))
        spng.set(qn('w:lineRule'), 'exact')
    pPr.append(spng)

def keep_with_next(para):
    pPr  = para._p.get_or_add_pPr()
    kwn  = OxmlElement('w:keepNext')
    pPr.append(kwn)

def page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pageBreakBefore')
    pPr.append(pb)

def add_run(para, text, bold=False, italic=False, size=None, colour=None, font=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:    run.font.size  = Pt(size)
    if colour:  run.font.color.rgb = colour
    if font:    run.font.name  = font
    return run

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


# ── Document helpers ──────────────────────────────────────────────────────────

def add_heading_para(doc, text, colour=DARK_BLUE, size=10, before=120, after=60):
    p = doc.add_paragraph()
    paragraph_spacing(p, before=before, after=after)
    # Underline via bottom border on paragraph
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '2')
    bot.set(qn('w:color'), hex_colour(MID_BLUE))
    pBdr.append(bot)
    pPr.append(pBdr)
    keep_with_next(p)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size  = Pt(size)
    r.font.color.rgb = colour
    r.font.name  = 'Calibri'
    return p

def add_body_para(doc, text, size=9.5, before=0, after=60, colour=BODY_TEXT, indent=None):
    p = doc.add_paragraph()
    paragraph_spacing(p, before=before, after=after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.color.rgb = colour
    r.font.name  = 'Calibri'
    return p

def add_bullet(doc, text, colour=MID_BLUE, size=9.5, indent_cm=0.4):
    p = doc.add_paragraph()
    paragraph_spacing(p, before=0, after=40)
    p.paragraph_format.left_indent  = Cm(indent_cm + 0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    # Arrow bullet
    run_bullet = p.add_run('▸  ')
    run_bullet.font.color.rgb = colour
    run_bullet.font.size = Pt(8)
    run_bullet.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.color.rgb = BODY_TEXT
    r.font.name  = 'Calibri'
    return p

def add_numbered(doc, num, text, size=9.5):
    p = doc.add_paragraph()
    paragraph_spacing(p, before=0, after=50)
    p.paragraph_format.left_indent  = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    rb = p.add_run(f'{num}.  ')
    rb.bold = True
    rb.font.color.rgb = WHITE
    rb.font.size = Pt(8)
    rb.font.name = 'Calibri'
    # Number circle via highlight trick — use shaded inline table instead;
    # simplest approach: coloured run
    rb2 = p.runs[0]
    # Shade the number using XML rPr highlight — not perfectly circular but professional
    rPr  = rb2._r.get_or_add_rPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour(MID_BLUE))
    rPr.append(shd)
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.color.rgb = BODY_TEXT
    r.font.name  = 'Calibri'
    return p


# ── Banner table (full-width coloured header) ─────────────────────────────────

def add_banner(doc, page=1):
    t = doc.add_table(rows=3, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)

    # Row 0: PSFGS brand  |  conference badge
    for cell in t.rows[0].cells:
        shade_cell(cell, DARK_BLUE)
        cell_spacing(cell, top=100, bottom=40, left=120, right=120)

    c_brand = t.rows[0].cells[0]
    c_conf  = t.rows[0].cells[1]
    t.columns[0].width = Cm(10)
    t.columns[1].width = Cm(7.5)

    # Brand
    p = c_brand.add_paragraph()
    paragraph_spacing(p, before=0, after=20)
    r = p.add_run('PSFGS')
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
    p2 = c_brand.add_paragraph()
    paragraph_spacing(p2, before=0, after=0)
    r2 = p2.add_run('PUBLIC SECTOR FINANCIAL GOVERNANCE SUITE')
    r2.font.size = Pt(7.5); r2.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r2.font.name = 'Calibri'

    # Conference badge
    c_conf.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p3 = c_conf.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_spacing(p3, before=0, after=20)
    r3 = p3.add_run('PRESENTED AT\n')
    r3.font.size = Pt(7); r3.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r3.font.name = 'Calibri'
    r4 = p3.add_run('SALGA National Members\' Assembly\n')
    r4.bold = True; r4.font.size = Pt(9); r4.font.color.rgb = WHITE; r4.font.name = 'Calibri'
    r5 = p3.add_run('July 2026  ·  Cape Town International Convention Centre')
    r5.font.size = Pt(7.5); r5.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r5.font.name = 'Calibri'

    # Row 1: separator line
    merged = t.rows[1].cells[0].merge(t.rows[1].cells[1])
    shade_cell(merged, DARK_BLUE)
    cell_spacing(merged, top=0, bottom=0, left=120, right=120)
    sep_p = merged.add_paragraph()
    paragraph_spacing(sep_p, before=0, after=0)
    pPr  = sep_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '0'); bot.set(qn('w:color'), '5B8EC4')
    pBdr.append(bot); pPr.append(pBdr)

    # Row 2: Main title
    merged2 = t.rows[2].cells[0].merge(t.rows[2].cells[1])
    shade_cell(merged2, DARK_BLUE)
    cell_spacing(merged2, top=60, bottom=100, left=120, right=120)

    if page == 1:
        p_title = merged2.add_paragraph()
        paragraph_spacing(p_title, before=0, after=30)
        r_t = p_title.add_run('Achieving Clean Audits Through Technology:')
        r_t.bold = True; r_t.font.size = Pt(14); r_t.font.color.rgb = WHITE; r_t.font.name = 'Calibri'
        p_sub = merged2.add_paragraph()
        paragraph_spacing(p_sub, before=0, after=40)
        r_s = p_sub.add_run('A Purpose-Built Governance Platform for South African Municipalities')
        r_s.bold = True; r_s.font.size = Pt(13); r_s.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r_s.font.name = 'Calibri'
        p_meta = merged2.add_paragraph()
        paragraph_spacing(p_meta, before=0, after=0)
        r_m = p_meta.add_run('Concept Note  ·  Submitted by Glance Management Technologies (Pty) Ltd  ·  Proprietary & Confidential')
        r_m.font.size = Pt(8); r_m.font.color.rgb = RGBColor(0x7F, 0x9F, 0xBF); r_m.font.name = 'Calibri'
    else:
        p_cont = merged2.add_paragraph()
        paragraph_spacing(p_cont, before=0, after=10)
        r_c = p_cont.add_run('PSFGS — Concept Note (continued)  ·  Glance Management Technologies (Pty) Ltd  ·  SALGA National Members\' Assembly · July 2026')
        r_c.font.size = Pt(8.5); r_c.font.color.rgb = WHITE; r_c.font.name = 'Calibri'

    doc.add_paragraph()  # spacing after banner


# ── Stat strip ────────────────────────────────────────────────────────────────

def add_stat_strip(doc):
    stats = [
        ('163',  'Municipalities with qualified, adverse or\ndisclaimer opinions (AGSA 2023/24)', RED),
        ('67%',  'Of all findings are repeat findings —\nsame issues raised year after year',    AMBER),
        ('R26bn','Irregular expenditure incurred by\nmunicipalities in 2023/24',                RED),
        ('90',   'Days — MFMA target resolution window\nper AGSA finding (Sec 131)',             DARK_BLUE),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)

    for i, (num, lbl, col) in enumerate(stats):
        cell = t.rows[0].cells[i]
        cell.width = Cm(4.37)
        shade_cell(cell, LIGHT_BLUE if col == DARK_BLUE else
                         RGBColor(0xFF, 0xF8, 0xE6) if col == AMBER else
                         RGBColor(0xFE, 0xF0, 0xEF))
        set_cell_border(cell, top=(hex_colour(col), 18, 'single'),
                               bottom=('FFFFFF', 0, 'none'),
                               left=('FFFFFF', 0, 'none'),
                               right=('FFFFFF', 0, 'none'))
        cell_spacing(cell, top=80, bottom=80, left=80, right=80)

        pn = cell.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_spacing(pn, before=0, after=20)
        rn = pn.add_run(num)
        rn.bold = True; rn.font.size = Pt(22); rn.font.color.rgb = col; rn.font.name = 'Calibri'

        pl = cell.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_spacing(pl, before=0, after=0)
        rl = pl.add_run(lbl)
        rl.font.size = Pt(7); rl.font.color.rgb = GREY_TEXT; rl.font.name = 'Calibri'

    doc.add_paragraph()


# ── Module grid (3×4) ─────────────────────────────────────────────────────────

def add_module_grid(doc):
    modules = [
        ('Audit Findings Tracker (AFT)',       'End-to-end finding lifecycle'),
        ('Budget vs Actual Monitor (BVM)',      'Real-time vote monitoring'),
        ('IFWE Register (IFW)',                 'Irregular expenditure tracking'),
        ('SCM Compliance Checker (SCC)',        'Procurement rule enforcement'),
        ('Section 71 Auto-Generator (S71)',     'NT monthly reports, automated'),
        ('Asset Verification System (AVS)',     'Physical asset reconciliation'),
        ('Consequence Management (CMT)',        'Case tracking & outcomes'),
        ('Delegation of Authority (DAR)',       'Approval matrix register'),
        ('Policy Compliance Manager (PCR)',     'Policy review & status'),
        ('Performance KPI Dashboard (PKD)',     'IDP targets vs actuals'),
        ('Executive Reporting Portal (ERP)',    'MM/Mayor governance view'),
        ('Entity Settings & Branding',          'Configurable per municipality'),
    ]
    t = doc.add_table(rows=4, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    idx = 0
    for row in t.rows:
        for cell in row.cells:
            if idx >= len(modules): break
            name, sub = modules[idx]; idx += 1
            shade_cell(cell, LIGHT_GREY)
            thick_left_border(cell, MID_BLUE)
            cell_spacing(cell, top=50, bottom=50, left=100, right=60)
            pn = cell.add_paragraph()
            paragraph_spacing(pn, before=0, after=10)
            rn = pn.add_run(name)
            rn.bold = True; rn.font.size = Pt(8); rn.font.color.rgb = DARK_BLUE; rn.font.name = 'Calibri'
            ps = cell.add_paragraph()
            paragraph_spacing(ps, before=0, after=0)
            rs = ps.add_run(sub)
            rs.font.size = Pt(7); rs.font.color.rgb = GREY_TEXT; rs.font.name = 'Calibri'
    doc.add_paragraph()


# ── Callout box ───────────────────────────────────────────────────────────────

def add_callout(doc, label, text, style='blue'):
    colours = {
        'blue':  (MID_BLUE,   LIGHT_BLUE,               RGBColor(0x1F, 0x38, 0x64)),
        'green': (DARK_GREEN, RGBColor(0xEE, 0xF8, 0xF2), RGBColor(0x1A, 0x4D, 0x2E)),
        'amber': (AMBER,      RGBColor(0xFF, 0xF8, 0xE6), RGBColor(0x5A, 0x3E, 0x00)),
    }
    border_col, bg_col, text_col = colours[style]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    cell = t.rows[0].cells[0]
    shade_cell(cell, bg_col)
    thick_left_border(cell, border_col)
    cell_spacing(cell, top=60, bottom=60, left=120, right=100)
    p = cell.add_paragraph()
    paragraph_spacing(p, before=0, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if label:
        rb = p.add_run(label + '  ')
        rb.bold = True; rb.font.size = Pt(9); rb.font.color.rgb = border_col; rb.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size = Pt(9); r.font.color.rgb = text_col; r.font.name = 'Calibri'
    doc.add_paragraph()


# ── Pricing table ─────────────────────────────────────────────────────────────

def add_pricing_table(doc):
    headers = ['Municipality Type', 'Count in SA', 'Annual Subscription Fee', 'Procurement Route']
    rows = [
        ('Metropolitan (Category A)', '8',   'R 180,000 per annum', 'Three Written Quotations'),
        ('District (Category C)',     '44',  'R  95,000 per annum', 'Three Written Quotations'),
        ('Local (Category B)',        '205', 'R  48,000 per annum', 'Three Written Quotations'),
        ('Full market (all 257)',     '257', 'R 16.7M per annum',   'Via SALGA Master Agreement'),
    ]
    t = doc.add_table(rows=5, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_cell(cell, DARK_BLUE)
        cell_spacing(cell, top=60, bottom=60, left=80, right=80)
        p = cell.add_paragraph()
        paragraph_spacing(p, before=0, after=0)
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

    # Data rows
    for ri, (mtype, count, fee, route) in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = LIGHT_BLUE if ri == 3 else (LIGHT_GREY if ri % 2 == 0 else WHITE)
        fee_col = DARK_BLUE if ri == 3 else DARK_GREEN
        for ci, val in enumerate([mtype, count, fee, route]):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            cell_spacing(cell, top=50, bottom=50, left=80, right=80)
            set_cell_border(cell, bottom=(hex_colour(RGBColor(0xE8, 0xEC, 0xF2)), 4, 'single'),
                                   top=('FFFFFF', 0, 'none'), left=('FFFFFF', 0, 'none'), right=('FFFFFF', 0, 'none'))
            p = cell.add_paragraph()
            paragraph_spacing(p, before=0, after=0)
            r = p.add_run(val)
            r.font.size  = Pt(8.5)
            r.font.name  = 'Calibri'
            r.font.color.rgb = fee_col if ci == 2 else (DARK_BLUE if ri == 3 else BODY_TEXT)
            if ci in (0, 2) or ri == 3:
                r.bold = (ci == 2 or ri == 3)

    doc.add_paragraph()


# ── Partner offer box ─────────────────────────────────────────────────────────

def add_partner_box(doc):
    items = [
        'Endorsed platform for all 257 member municipalities',
        'SALGA co-branding prominently on the platform portal',
        'Quarterly aggregated governance analytics for SALGA reporting',
        '10% of net revenue contributed to the SALGA Capacity Fund',
        'Free deployment for 5 SALGA-nominated pilot municipalities',
        'Annual Clean Audit Progress Report delivered at SALGA AGM',
    ]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    cell = t.rows[0].cells[0]
    shade_cell(cell, LIGHT_BLUE)
    set_cell_border(cell,
        top=(hex_colour(MID_BLUE), 8, 'single'), bottom=(hex_colour(MID_BLUE), 8, 'single'),
        left=(hex_colour(MID_BLUE), 8, 'single'), right=(hex_colour(MID_BLUE), 8, 'single'))
    cell_spacing(cell, top=80, bottom=80, left=120, right=120)

    pt = cell.add_paragraph()
    paragraph_spacing(pt, before=0, after=60)
    rt = pt.add_run('What Glance Management Technologies Offers SALGA')
    rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = DARK_BLUE; rt.font.name = 'Calibri'

    for item in items:
        pi = cell.add_paragraph()
        paragraph_spacing(pi, before=0, after=30)
        pi.paragraph_format.left_indent  = Cm(0.35)
        pi.paragraph_format.first_line_indent = Cm(-0.35)
        rc = pi.add_run('✓  ')
        rc.bold = True; rc.font.size = Pt(9); rc.font.color.rgb = DARK_GREEN; rc.font.name = 'Calibri'
        ri = pi.add_run(item)
        ri.font.size = Pt(9); ri.font.color.rgb = BODY_TEXT; ri.font.name = 'Calibri'

    doc.add_paragraph()


# ── CTA footer banner ─────────────────────────────────────────────────────────

def add_cta_banner(doc):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    t.columns[0].width = Cm(10.5)
    t.columns[1].width = Cm(7)

    for cell in t.rows[0].cells:
        shade_cell(cell, DARK_BLUE)
        cell_spacing(cell, top=100, bottom=100, left=140, right=140)

    left  = t.rows[0].cells[0]
    right = t.rows[0].cells[1]

    # Left — CTA text
    ph = left.add_paragraph()
    paragraph_spacing(ph, before=0, after=30)
    rh = ph.add_run('Ready to pilot PSFGS at your municipality?')
    rh.bold = True; rh.font.size = Pt(11); rh.font.color.rgb = WHITE; rh.font.name = 'Calibri'

    pb = left.add_paragraph()
    paragraph_spacing(pb, before=0, after=0)
    rb = pb.add_run(
        'Request a live demonstration at the SALGA conference stand or schedule a '
        'remote walkthrough with your Internal Audit and CFO teams. '
        'First 5 SALGA-nominated municipalities onboard free of charge.')
    rb.font.size = Pt(8.5); rb.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); rb.font.name = 'Calibri'

    # Right — contact
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for line in [
        ('Glance Management Technologies (Pty) Ltd', True,  9.5, WHITE),
        ('info@glancemanagement.co.za',               False, 8.5, RGBColor(0xAE, 0xD6, 0xF1)),
        ('www.psfgs.co.za',                           False, 8.5, RGBColor(0xAE, 0xD6, 0xF1)),
        ('+27 (0) 11 000 0000',                       False, 8.5, RGBColor(0xAE, 0xD6, 0xF1)),
        ('CSD Registered  ·  B-BBEE Level 1',         False, 7.5, RGBColor(0x7F, 0x9F, 0xBF)),
    ]:
        text, bold, size, col = line
        pc = right.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_spacing(pc, before=0, after=20)
        rc = pc.add_run(text)
        rc.bold = bold; rc.font.size = Pt(size); rc.font.color.rgb = col; rc.font.name = 'Calibri'


# ── Footer bar (copyright strip) ─────────────────────────────────────────────

def add_footer_strip(doc):
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)

    data = [
        ('PSFGS — Public Sector Financial Governance Suite\nBuilt for SA municipalities · MFMA-native · Cloud SaaS',
         WD_ALIGN_PARAGRAPH.LEFT),
        ('© 2026 Glance Management Technologies (Pty) Ltd\nAll rights reserved · Proprietary & Confidential',
         WD_ALIGN_PARAGRAPH.CENTER),
        ('SALGA National Members\' Assembly\nJuly 2026 · Cape Town',
         WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for i, (text, align) in enumerate(data):
        cell = t.rows[0].cells[i]
        shade_cell(cell, DARK_BLUE)
        cell_spacing(cell, top=60, bottom=60, left=100, right=100)
        p = cell.add_paragraph()
        p.alignment = align
        paragraph_spacing(p, before=0, after=0)
        r = p.add_run(text)
        r.font.size = Pt(7); r.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r.font.name = 'Calibri'


# ═════════════════════════════════════════════════════════════════════════════
#  MAIN BUILD
# ═════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Default style ─────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9.5)

    # ══════════════════════════  PAGE 1  ═════════════════════════════════════

    add_banner(doc, page=1)
    add_stat_strip(doc)

    # Problem
    add_heading_para(doc, 'The Problem: Governance Failures Are Systemic and Recurring')
    p1 = add_body_para(doc,
        'South African municipalities face a compounding governance crisis. Despite billions invested in audit support, '
        'consultants, and training, the number of municipalities achieving clean audits has stagnated below 30 for six '
        'consecutive years. The root cause is not a lack of awareness — it is a lack of structured, real-time tooling '
        'to track, remediate, and evidence corrective action between audit cycles.')
    p2 = add_body_para(doc,
        'Accounting Officers, CFOs, and Directors operate across disconnected spreadsheets, email chains, and manual '
        'registers. Findings raised in October are forgotten by March. Evidence submitted to Internal Audit is lost or '
        'unverifiable. Section 71 reports are compiled at month-end in a panic. Consequence management cases stall for '
        'years without structured tracking. The AGSA raises the same findings. Audit opinions deteriorate further.')
    add_callout(doc,
        'The cost of non-compliance:',
        'A municipality carrying a Qualified opinion forfeits credibility with investors, development finance institutions, '
        'and the National Treasury — triggering Section 139 interventions, withheld equitable share allocations, and '
        'reputational damage that directly affects service delivery for millions of residents.',
        style='amber')

    # Solution
    add_heading_para(doc, 'The Solution: PSFGS — Governance Infrastructure for Every Municipality')
    add_body_para(doc,
        'The Public Sector Financial Governance Suite (PSFGS) is a South African-built, MFMA-native, cloud-based '
        'platform that digitises and automates the full spectrum of municipal financial governance — from the moment '
        'an audit finding is raised to the day it is formally closed and evidenced to the AGSA. The platform comprises '
        'twelve integrated modules:')
    add_module_grid(doc)

    # Two-column: How it works | Who uses it
    t2 = doc.add_table(rows=1, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t2)
    t2.columns[0].width = Cm(8.7)
    t2.columns[1].width = Cm(8.7)

    def fill_col(cell, heading, bullets):
        cell_spacing(cell, top=0, bottom=0, left=0, right=80)
        ph = cell.add_paragraph()
        paragraph_spacing(ph, before=60, after=60)
        pPr  = ph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
        bot.set(qn('w:space'), '2');    bot.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr.append(bot); pPr.append(pBdr)
        rh = ph.add_run(heading.upper())
        rh.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = DARK_BLUE; rh.font.name = 'Calibri'
        for b_label, b_text in bullets:
            pb = cell.add_paragraph()
            paragraph_spacing(pb, before=0, after=35)
            pb.paragraph_format.left_indent = Cm(0.35)
            pb.paragraph_format.first_line_indent = Cm(-0.35)
            rc = pb.add_run('▸  ')
            rc.font.size = Pt(8); rc.font.color.rgb = MID_BLUE; rc.font.name = 'Calibri'
            if b_label:
                rl = pb.add_run(b_label + '  ')
                rl.bold = True; rl.font.size = Pt(9); rl.font.color.rgb = DARK_BLUE; rl.font.name = 'Calibri'
            rt = pb.add_run(b_text)
            rt.font.size = Pt(9); rt.font.color.rgb = BODY_TEXT; rt.font.name = 'Calibri'

    fill_col(t2.rows[0].cells[0], 'How It Works', [
        ('Cloud-hosted SaaS —', 'no server or hardware procurement required'),
        ('Multi-tenant —', 'each municipality\'s data is fully isolated and secure'),
        ('Role-based access —', 'Auditor, Director, CFO and MM each see only relevant data'),
        ('MFMA workflow-aligned —', 'finding stages mirror AGSA protocols exactly'),
        ('AI-assisted analysis —', 'identifies trends, repeat findings and at-risk budget votes'),
        ('Configurable branding —', 'entity logo, name and leadership details per municipality'),
    ])
    fill_col(t2.rows[0].cells[1], 'Who Uses It', [
        ('Internal Audit —', 'captures and tracks every finding with supporting evidence'),
        ('City Manager / MM —', 'assigns findings, monitors remediation against due dates'),
        ('CFO —', 'approves closures, reviews Section 71 reports before submission'),
        ('Directorate Managers —', 'implements corrective actions with structured deadlines'),
        ('Council / MPAC —', 'oversight dashboard with automatic escalation notifications'),
        ('National Treasury / AGSA —', 'structured evidence package ready at year-end'),
    ])

    # ══════════════════════════  PAGE 2  ═════════════════════════════════════

    doc.add_page_break()
    add_banner(doc, page=2)

    # Two-column layout for page 2
    tp2 = doc.add_table(rows=1, cols=2)
    tp2.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tp2)
    tp2.columns[0].width = Cm(8.7)
    tp2.columns[1].width = Cm(8.7)
    left2  = tp2.rows[0].cells[0]
    right2 = tp2.rows[0].cells[1]
    cell_spacing(left2,  top=0, bottom=0, left=0, right=80)
    cell_spacing(right2, top=0, bottom=0, left=80, right=0)

    def add_section_to_cell(cell, heading, content_fn):
        ph = cell.add_paragraph()
        paragraph_spacing(ph, before=0, after=60)
        pPr  = ph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
        bot.set(qn('w:space'), '2');    bot.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr.append(bot); pPr.append(pBdr)
        rh = ph.add_run(heading.upper())
        rh.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = DARK_BLUE; rh.font.name = 'Calibri'
        content_fn(cell)

    # ── LEFT column content ───────────────────────────────────────────────────

    def left_col_content(cell):
        # Why SALGA
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_spacing(p1, before=0, after=60)
        r1 = p1.add_run(
            'SALGA\'s mandate is to strengthen local government capacity. PSFGS directly advances three of '
            'SALGA\'s 2024–2029 strategic priorities: ')
        r1.font.size = Pt(9.5); r1.font.color.rgb = BODY_TEXT; r1.font.name = 'Calibri'
        rb = p1.add_run('financial sustainability, good governance, and audit outcome improvement.')
        rb.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = DARK_BLUE; rb.font.name = 'Calibri'

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_spacing(p2, before=0, after=100)
        r2 = p2.add_run(
            'A SALGA-endorsed master agreement removes the need for each of the 257 municipalities to run individual '
            'procurement processes. Municipalities access the platform under a cooperative procurement arrangement — '
            'legally sound under MFMA SCM Regulations — cutting procurement time from six months to two weeks.')
        r2.font.size = Pt(9.5); r2.font.color.rgb = BODY_TEXT; r2.font.name = 'Calibri'

        # Procurement section
        ph2 = cell.add_paragraph()
        paragraph_spacing(ph2, before=0, after=60)
        pPr  = ph2._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
        bot.set(qn('w:space'), '2');    bot.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr.append(bot); pPr.append(pBdr)
        rh2 = ph2.add_run('PROCUREMENT — NO TENDER REQUIRED')
        rh2.bold = True; rh2.font.size = Pt(9); rh2.font.color.rgb = DARK_BLUE; rh2.font.name = 'Calibri'

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_spacing(p3, before=0, after=80)
        r3 = p3.add_run(
            'PSFGS is priced deliberately within the R200,000 threshold that permits procurement by written quotation '
            'under MFMA SCM Regulations. Individual municipalities require no public tender advertisement, no CIDB '
            'registration, and no council resolution — only three comparative quotes and a CFO sign-off under '
            'existing delegations of authority.')
        r3.font.size = Pt(9.5); r3.font.color.rgb = BODY_TEXT; r3.font.name = 'Calibri'

        # Callout in cell
        tc_inner = cell.add_table(rows=1, cols=1)
        no_borders(tc_inner)
        ci = tc_inner.rows[0].cells[0]
        shade_cell(ci, LIGHT_BLUE)
        thick_left_border(ci, MID_BLUE)
        cell_spacing(ci, top=60, bottom=60, left=100, right=80)
        pc = ci.add_paragraph()
        paragraph_spacing(pc, before=0, after=0)
        pc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rb2 = pc.add_run('SALGA Master Agreement route:  ')
        rb2.bold = True; rb2.font.size = Pt(8.5); rb2.font.color.rgb = MID_BLUE; rb2.font.name = 'Calibri'
        rc2 = pc.add_run(
            'One cooperative procurement event establishes the framework. All member municipalities subscribe directly. '
            'Estimated time-to-activation per municipality: 10–15 working days.')
        rc2.font.size = Pt(8.5); rc2.font.color.rgb = RGBColor(0x1F, 0x38, 0x64); rc2.font.name = 'Calibri'

        cell.add_paragraph()

        # Pricing heading
        ph3 = cell.add_paragraph()
        paragraph_spacing(ph3, before=60, after=60)
        pPr3  = ph3._p.get_or_add_pPr()
        pBdr3 = OxmlElement('w:pBdr')
        bot3  = OxmlElement('w:bottom')
        bot3.set(qn('w:val'), 'single'); bot3.set(qn('w:sz'), '12')
        bot3.set(qn('w:space'), '2');    bot3.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr3.append(bot3); pPr3.append(pBdr3)
        rh3 = ph3.add_run('PRICING — AFFORDABLE AT EVERY TIER')
        rh3.bold = True; rh3.font.size = Pt(9); rh3.font.color.rgb = DARK_BLUE; rh3.font.name = 'Calibri'

        # Pricing table inside cell
        pt = cell.add_table(rows=5, cols=3)
        no_borders(pt)
        pthdrs = ['Municipality Type', 'Annual Fee', 'Route']
        pt_data = [
            ('Metropolitan (Cat A) — 8 municipalities',  'R 180,000 / yr', '3 Quotations'),
            ('District (Cat C) — 44 municipalities',     'R  95,000 / yr', '3 Quotations'),
            ('Local (Cat B) — 205 municipalities',       'R  48,000 / yr', '3 Quotations'),
            ('Full market — 257 municipalities',         'R 16.7M / yr',   'SALGA Agreement'),
        ]
        for ci2, h in enumerate(pthdrs):
            c = pt.rows[0].cells[ci2]
            shade_cell(c, DARK_BLUE)
            cell_spacing(c, top=40, bottom=40, left=60, right=60)
            pp = c.add_paragraph()
            paragraph_spacing(pp, before=0, after=0)
            rr = pp.add_run(h)
            rr.bold = True; rr.font.size = Pt(7.5); rr.font.color.rgb = WHITE; rr.font.name = 'Calibri'
        for ri2, (mtype2, fee2, route2) in enumerate(pt_data):
            bg2 = LIGHT_BLUE if ri2 == 3 else (LIGHT_GREY if ri2 % 2 == 0 else WHITE)
            row2 = pt.rows[ri2 + 1]
            for ci3, val2 in enumerate([mtype2, fee2, route2]):
                c2 = row2.cells[ci3]
                shade_cell(c2, bg2)
                cell_spacing(c2, top=35, bottom=35, left=60, right=60)
                set_cell_border(c2, bottom=(hex_colour(RGBColor(0xE8, 0xEC, 0xF2)), 4, 'single'),
                                    top=('FFFFFF', 0, 'none'), left=('FFFFFF', 0, 'none'), right=('FFFFFF', 0, 'none'))
                pp2 = c2.add_paragraph()
                paragraph_spacing(pp2, before=0, after=0)
                rr2 = pp2.add_run(val2)
                rr2.font.size = Pt(8)
                rr2.font.name = 'Calibri'
                rr2.font.color.rgb = DARK_BLUE if ri2 == 3 else (DARK_GREEN if ci3 == 1 else BODY_TEXT)
                rr2.bold = (ri2 == 3 or ci3 == 1)

        pnote = cell.add_paragraph()
        paragraph_spacing(pnote, before=30, after=0)
        rnote = pnote.add_run('All pricing excludes VAT. Optional implementation support (R15,000 once-off) available for first-time onboarding.')
        rnote.font.size = Pt(7); rnote.font.color.rgb = GREY_TEXT; rnote.font.name = 'Calibri'

    add_section_to_cell(left2, 'Why Now? The SALGA Opportunity', left_col_content)

    # ── RIGHT column content ──────────────────────────────────────────────────

    def right_col_content(cell):
        # Partner box heading
        ph_r = cell.add_paragraph()
        paragraph_spacing(ph_r, before=0, after=60)
        pPr_r  = ph_r._p.get_or_add_pPr()
        pBdr_r = OxmlElement('w:pBdr')
        bot_r  = OxmlElement('w:bottom')
        bot_r.set(qn('w:val'), 'single'); bot_r.set(qn('w:sz'), '12')
        bot_r.set(qn('w:space'), '2');    bot_r.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr_r.append(bot_r); pPr_r.append(pBdr_r)
        rh_r = ph_r.add_run('PROPOSED SALGA PARTNERSHIP MODEL')
        rh_r.bold = True; rh_r.font.size = Pt(9); rh_r.font.color.rgb = DARK_BLUE; rh_r.font.name = 'Calibri'

        # Partner offers
        offers = [
            'Endorsed platform accessible to all 257 member municipalities',
            'SALGA co-branding prominently displayed on the platform portal',
            'Quarterly aggregated governance analytics reports for SALGA',
            '10% of net annual revenue contributed to SALGA Capacity Fund',
            'Free deployment for 5 SALGA-nominated pilot municipalities',
            'Annual Clean Audit Progress Report delivered at SALGA AGM',
        ]
        pt_inner = cell.add_table(rows=1, cols=1)
        no_borders(pt_inner)
        ci_i = pt_inner.rows[0].cells[0]
        shade_cell(ci_i, LIGHT_BLUE)
        set_cell_border(ci_i,
            top=(hex_colour(MID_BLUE), 6, 'single'), bottom=(hex_colour(MID_BLUE), 6, 'single'),
            left=(hex_colour(MID_BLUE), 6, 'single'), right=(hex_colour(MID_BLUE), 6, 'single'))
        cell_spacing(ci_i, top=60, bottom=60, left=100, right=80)
        ptitle = ci_i.add_paragraph()
        paragraph_spacing(ptitle, before=0, after=50)
        rt_i = ptitle.add_run('What GMT Offers SALGA')
        rt_i.bold = True; rt_i.font.size = Pt(9.5); rt_i.font.color.rgb = DARK_BLUE; rt_i.font.name = 'Calibri'
        for offer in offers:
            po = ci_i.add_paragraph()
            paragraph_spacing(po, before=0, after=25)
            po.paragraph_format.left_indent = Cm(0.35)
            po.paragraph_format.first_line_indent = Cm(-0.35)
            rc_i = po.add_run('✓  ')
            rc_i.bold = True; rc_i.font.size = Pt(9); rc_i.font.color.rgb = DARK_GREEN; rc_i.font.name = 'Calibri'
            ri_i = po.add_run(offer)
            ri_i.font.size = Pt(8.5); ri_i.font.color.rgb = BODY_TEXT; ri_i.font.name = 'Calibri'

        cell.add_paragraph()

        # Rollout heading
        ph_roll = cell.add_paragraph()
        paragraph_spacing(ph_roll, before=0, after=60)
        pPr_roll  = ph_roll._p.get_or_add_pPr()
        pBdr_roll = OxmlElement('w:pBdr')
        bot_roll  = OxmlElement('w:bottom')
        bot_roll.set(qn('w:val'), 'single'); bot_roll.set(qn('w:sz'), '12')
        bot_roll.set(qn('w:space'), '2');    bot_roll.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr_roll.append(bot_roll); pPr_roll.append(pBdr_roll)
        rh_roll = ph_roll.add_run('IMPLEMENTATION — PILOT TO NATIONAL ROLLOUT')
        rh_roll.bold = True; rh_roll.font.size = Pt(9); rh_roll.font.color.rgb = DARK_BLUE; rh_roll.font.name = 'Calibri'

        steps = [
            ('Q3 2026 (Aug–Sep)', 'Onboard 5 SALGA-nominated pilot municipalities across 3 provinces at zero cost to pilots.'),
            ('Q4 2026 (Oct–Nov)', 'Present pilot outcomes at COGTA Municipal Finance Indaba. Initiate master agreement finalisation.'),
            ('Q1 2027 (Jan–Mar)', 'Open subscription to all 257 municipalities. Target 25 active clients by 31 March.'),
            ('FY 2027/28',        'Scale to 80+ municipalities. Integrate with NT eSourcing portal and AGSA management letter system.'),
            ('FY 2028/29',        'Full national coverage. Introduce predictive audit risk scoring and NT real-time reporting API.'),
        ]
        for i, (period, desc) in enumerate(steps):
            ps = cell.add_paragraph()
            paragraph_spacing(ps, before=0, after=40)
            ps.paragraph_format.left_indent = Cm(0.8)
            ps.paragraph_format.first_line_indent = Cm(-0.8)
            rnum = ps.add_run(f'{i+1}  ')
            rnum.bold = True; rnum.font.size = Pt(7.5); rnum.font.color.rgb = WHITE; rnum.font.name = 'Calibri'
            rPr_n = rnum._r.get_or_add_rPr()
            shd_n = OxmlElement('w:shd')
            shd_n.set(qn('w:val'), 'clear'); shd_n.set(qn('w:color'), 'auto')
            shd_n.set(qn('w:fill'), hex_colour(MID_BLUE))
            rPr_n.append(shd_n)
            rp = ps.add_run(f'{period}:  ')
            rp.bold = True; rp.font.size = Pt(9); rp.font.color.rgb = DARK_BLUE; rp.font.name = 'Calibri'
            rd = ps.add_run(desc)
            rd.font.size = Pt(9); rd.font.color.rgb = BODY_TEXT; rd.font.name = 'Calibri'

        cell.add_paragraph()

        # Regulatory heading
        ph_reg = cell.add_paragraph()
        paragraph_spacing(ph_reg, before=0, after=60)
        pPr_reg  = ph_reg._p.get_or_add_pPr()
        pBdr_reg = OxmlElement('w:pBdr')
        bot_reg  = OxmlElement('w:bottom')
        bot_reg.set(qn('w:val'), 'single'); bot_reg.set(qn('w:sz'), '12')
        bot_reg.set(qn('w:space'), '2');    bot_reg.set(qn('w:color'), hex_colour(MID_BLUE))
        pBdr_reg.append(bot_reg); pPr_reg.append(pBdr_reg)
        rh_reg = ph_reg.add_run('REGULATORY ALIGNMENT & TRUST SIGNALS')
        rh_reg.bold = True; rh_reg.font.size = Pt(9); rh_reg.font.color.rgb = DARK_BLUE; rh_reg.font.name = 'Calibri'

        reg_bullets = [
            ('MFMA Sections 32, 62, 71, 72, 121 & 131 —', 'all compliance requirements embedded in platform workflows'),
            ('AGSA Finding Protocol —', 'lifecycle stages mirror AG management letter structure exactly'),
            ('NT Circular 88 & Instruction Note 5 —', 'budget monitoring aligned to National Treasury reporting formats'),
            ('POPIA Compliant —', 'data hosted in South Africa; role-based access; full immutable audit trail'),
            ('CSD Registered Supplier —', 'Central Supplier Database active; B-BBEE Level 1 certification'),
        ]
        for b_label, b_text in reg_bullets:
            pb = cell.add_paragraph()
            paragraph_spacing(pb, before=0, after=30)
            pb.paragraph_format.left_indent = Cm(0.35)
            pb.paragraph_format.first_line_indent = Cm(-0.35)
            rc = pb.add_run('▸  ')
            rc.font.size = Pt(8); rc.font.color.rgb = MID_BLUE; rc.font.name = 'Calibri'
            rl = pb.add_run(b_label + '  ')
            rl.bold = True; rl.font.size = Pt(9); rl.font.color.rgb = DARK_BLUE; rl.font.name = 'Calibri'
            rt = pb.add_run(b_text)
            rt.font.size = Pt(9); rt.font.color.rgb = BODY_TEXT; rt.font.name = 'Calibri'

        # Reference client callout inside cell
        cell.add_paragraph()
        tc_ref = cell.add_table(rows=1, cols=1)
        no_borders(tc_ref)
        ci_ref = tc_ref.rows[0].cells[0]
        shade_cell(ci_ref, RGBColor(0xEE, 0xF8, 0xF2))
        thick_left_border(ci_ref, DARK_GREEN)
        cell_spacing(ci_ref, top=60, bottom=60, left=100, right=80)
        pc_ref = ci_ref.add_paragraph()
        paragraph_spacing(pc_ref, before=0, after=0)
        pc_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rb_ref = pc_ref.add_run('Reference Client:  ')
        rb_ref.bold = True; rb_ref.font.size = Pt(8.5); rb_ref.font.color.rgb = DARK_GREEN; rb_ref.font.name = 'Calibri'
        rc_ref = pc_ref.add_run(
            'Buffalo City Metropolitan Municipality (BCMM) — fully operational deployment across 10 directorates '
            'with 35 role-based user accounts. Live audit finding workflow, Section 71 reporting, SCM compliance '
            'checking, and IFWE register in production as of April 2026.')
        rc_ref.font.size = Pt(8.5); rc_ref.font.color.rgb = RGBColor(0x1A, 0x4D, 0x2E); rc_ref.font.name = 'Calibri'

    right2.add_paragraph()  # top padding
    right_col_content(right2)

    # CTA and footer
    doc.add_paragraph()
    add_cta_banner(doc)
    add_footer_strip(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document saved: {OUTPUT_PATH}")


if __name__ == '__main__':
    build()
