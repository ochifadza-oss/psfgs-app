"""PSFGS Training Manual Generator — python-docx"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "PSFGS_Training_Manual.docx")

# Colours
DB  = RGBColor(0x1F,0x38,0x64)
MB  = RGBColor(0x2E,0x75,0xB6)
LB  = RGBColor(0xEE,0xF5,0xFC)
DG  = RGBColor(0x1A,0x6B,0x3C)
AM  = RGBColor(0xD4,0x86,0x1A)
WH  = RGBColor(0xFF,0xFF,0xFF)
BD  = RGBColor(0x2C,0x3E,0x50)
GR  = RGBColor(0x5A,0x64,0x78)
LLB = RGBColor(0xAE,0xD6,0xF1)
LG  = RGBColor(0xF4,0xF6,0xF9)
AM2 = RGBColor(0xFF,0xF8,0xE6)
DG2 = RGBColor(0xEE,0xF8,0xF2)

def hx(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def shade(cell, rgb):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hx(rgb))
    pr.append(s)

def no_border(cell):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    b=OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        e=OxmlElement(f'w:{side}'); e.set(qn('w:val'),'none')
        b.append(e)
    pr.append(b)

def cell_mar(cell,t=60,b=60,l=100,r=100):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    m=OxmlElement('w:tcMar')
    for s,v in [('top',t),('bottom',b),('left',l),('right',r)]:
        e=OxmlElement(f'w:{s}'); e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa')
        m.append(e)
    pr.append(m)

def thick_left(cell, rgb):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    b=OxmlElement('w:tcBorders')
    for side in ('top','bottom','right'):
        e=OxmlElement(f'w:{side}'); e.set(qn('w:val'),'none'); b.append(e)
    e=OxmlElement('w:left'); e.set(qn('w:val'),'single')
    e.set(qn('w:sz'),'24'); e.set(qn('w:color'),hx(rgb)); b.append(e)
    pr.append(b)

def sp(p,bef=0,aft=0):
    pr=p._p.get_or_add_pPr()
    s=OxmlElement('w:spacing')
    s.set(qn('w:before'),str(bef)); s.set(qn('w:after'),str(aft))
    pr.append(s)

def run(p,txt,bold=False,size=9.5,col=None,font='Calibri',italic=False):
    r=p.add_run(txt); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.name=font
    if col: r.font.color.rgb=col
    return r

# ── helpers ──────────────────────────────────────────────────────────────────
def para(doc,txt='',bold=False,size=9.5,col=None,align=WD_ALIGN_PARAGRAPH.JUSTIFY,bef=0,aft=60,indent=None):
    p=doc.add_paragraph(); p.alignment=align; sp(p,bef,aft)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    if txt: run(p,txt,bold,size,col or BD)
    return p

def heading(doc,txt,size=9.5,bef=120,aft=60):
    p=doc.add_paragraph(); sp(p,bef,aft)
    pr=p._p.get_or_add_pPr()
    bd=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12')
    bot.set(qn('w:space'),'2'); bot.set(qn('w:color'),hx(MB))
    bd.append(bot); pr.append(bd)
    run(p,txt.upper(),True,size,DB)
    return p

def chapter_banner(doc,txt):
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0]; shade(c,DB); no_border(c); cell_mar(c,100,100,140,140)
    p=c.add_paragraph(); sp(p,0,0)
    run(p,txt,True,13,WH)
    doc.add_paragraph(); return t

def bullet(doc,txt,bold_prefix=None,bef=0,aft=40,indent=0.4):
    p=doc.add_paragraph(); sp(p,bef,aft)
    p.paragraph_format.left_indent=Cm(indent+0.35)
    p.paragraph_format.first_line_indent=Cm(-0.35)
    r=p.add_run('  '); r.font.color.rgb=MB; r.font.size=Pt(8); r.font.name='Calibri'
    if bold_prefix:
        run(p,bold_prefix+'  ',True,9.5,DB)
    run(p,txt,False,9.5,BD)
    return p

def step(doc,num,txt):
    p=doc.add_paragraph(); sp(p,0,50)
    p.paragraph_format.left_indent=Cm(0.8)
    p.paragraph_format.first_line_indent=Cm(-0.8)
    r=p.add_run(f'Step {num}:  '); r.bold=True; r.font.color.rgb=MB; r.font.size=Pt(9.5); r.font.name='Calibri'
    run(p,txt,False,9.5,BD)

def callout(doc,label,txt,bg=AM2,border=AM):
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0]; shade(c,bg); thick_left(c,border); no_border(c)
    cell_mar(c,80,80,120,100)
    p=c.add_paragraph(); sp(p,0,0); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if label: run(p,label+'  ',True,9,border)
    run(p,txt,False,9,BD)
    doc.add_paragraph()

def error_table(doc,rows):
    t=doc.add_table(len(rows)+1,3)
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    widths=[Cm(5.5),Cm(5.5),Cm(6.5)]
    hdrs=['Error Message / Symptom','Likely Cause','Resolution']
    for i,(h,w) in enumerate(zip(hdrs,widths)):
        c=t.rows[0].cells[i]; c.width=w
        shade(c,DB); no_border(c); cell_mar(c,60,60,80,80)
        p=c.add_paragraph(); sp(p,0,0); run(p,h,True,8,WH)
    for ri,(e,ca,res) in enumerate(rows):
        for ci,(v,w) in enumerate(zip([e,ca,res],widths)):
            c=t.rows[ri+1].cells[ci]; c.width=w
            shade(c,LG if ri%2==0 else WH); no_border(c); cell_mar(c,50,50,80,80)
            p=c.add_paragraph(); sp(p,0,0)
            run(p,v,False,8.5,BD)
    doc.add_paragraph()

def section_label(doc,txt,bef=100,aft=40):
    p=doc.add_paragraph(); sp(p,bef,aft)
    run(p,txt,True,9.5,MB)

# ── module builder ────────────────────────────────────────────────────────────
def module_chapter(doc, ch_num, code, name,
                   purpose_paras, legislation,
                   users,       # list of (role, action)
                   features,    # list of (name, desc)
                   steps,       # list of (num, text)
                   tip_label, tip_text,
                   errors):     # list of (error, cause, fix)
    doc.add_page_break()
    chapter_banner(doc, f'CHAPTER {ch_num}: {name} ({code})')

    section_label(doc, f'{ch_num}.1  Purpose and Legislative Basis')
    for p in purpose_paras: para(doc, p)
    para(doc, f'Legislative basis: {legislation}', False, 9, GR)

    section_label(doc, f'{ch_num}.2  Who Uses This Module')
    for role, action in users:
        bullet(doc, action, bold_prefix=role)

    section_label(doc, f'{ch_num}.3  Key Features')
    for fname, fdesc in features:
        bullet(doc, fdesc, bold_prefix=fname)

    section_label(doc, f'{ch_num}.4  Step-by-Step Workflow')
    for n, txt in steps: step(doc, n, txt)

    section_label(doc, f'{ch_num}.5  Important Notes and Tips')
    callout(doc, tip_label, tip_text)

    section_label(doc, f'{ch_num}.6  Common Errors and How to Fix')
    error_table(doc, errors)

# ═════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    doc.styles['Normal'].font.name='Calibri'
    doc.styles['Normal'].font.size=Pt(9.5)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0]; shade(c,DB); no_border(c); cell_mar(c,500,500,200,200)

    for txt,sz,col,ital,bef,aft in [
        ('PSFGS',48,WH,False,0,40),
        ('PUBLIC SECTOR FINANCIAL GOVERNANCE SUITE',11,LLB,False,0,200),
        ('USER TRAINING MANUAL',28,WH,False,100,40),
        ('Complete Guide to All 13 Modules',14,LLB,True,0,160),
    ]:
        p=c.add_paragraph(); sp(p,bef,aft); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,txt,sz>=16,sz,col,italic=ital)

    p=c.add_paragraph(); sp(p,0,30); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('_'*60); r.font.color.rgb=WH; r.font.size=Pt(9)

    for txt,sz,col in [
        ('Version 1.0  |  April 2026  |  For Official Use Only',9,LLB),
        ('Prepared by Glance Management Technologies (Pty) Ltd',9,LLB),
        ('MFMA-Compliant  |  Cloud SaaS  |  Multi-Tenant Architecture',8,RGBColor(0x7F,0x9F,0xBF)),
    ]:
        p=c.add_paragraph(); sp(p,0,20); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,txt,False,sz,col)

    # ── PAGE 2: DOC INFO + TOC ────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc,'Document Information',10,0,80)
    di_rows=[
        ('Document Title','PSFGS User Training Manual — Complete Guide to All 13 Modules'),
        ('Version','1.0'),('Date','April 2026'),
        ('Prepared By','Glance Management Technologies (Pty) Ltd'),
        ('Reviewed By','PSFGS Product Team'),
        ('Classification','For Official Use — Municipality Staff Only'),
        ('Distribution','All PSFGS system users; System Administrators; Municipal Management'),
    ]
    dt=doc.add_table(len(di_rows),2); dt.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,(label,val) in enumerate(di_rows):
        bg=LB if ri%2==0 else WH
        for ci,(v,w) in enumerate([(label,Cm(5)),(val,Cm(12.5))]):
            c=dt.rows[ri].cells[ci]; c.width=w; shade(c,bg); no_border(c); cell_mar(c,50,50,80,80)
            p=c.add_paragraph(); sp(p,0,0)
            run(p,v,ci==0,8.5,DB if ci==0 else BD)
    doc.add_paragraph()

    heading(doc,'Table of Contents',10,60,60)
    toc_items=[
        ('1','Introduction to PSFGS','3'),
        ('  1.1','About the System','3'),('  1.2','System Architecture','3'),
        ('  1.3','Logging In','3'),('  1.4','Navigation','4'),('  1.5','User Roles','4'),
        ('2','Executive Dashboard (DASH)','5'),
        ('3','Audit Findings Tracker (AFT)','7'),
        ('4','Asset Verification System (AVS)','10'),
        ('5','Budget vs Actual Monitor (BVM)','12'),
        ('6','IFWE Register (IFW)','14'),
        ('7','SCM Compliance Checker (SCC)','17'),
        ('8','Consequence Management Tracker (CMT)','19'),
        ('9','Delegation of Authority Register (DAR)','22'),
        ('10','Policy Compliance Manager (PCR)','24'),
        ('11','Section 71 Auto-Generator (S71)','27'),
        ('12','Performance KPI Dashboard (PKD)','29'),
        ('13','Executive Reporting Portal (ERP)','32'),
        ('14','Entity Settings','34'),
        ('Appendix A','User Roles and Permissions Matrix','36'),
        ('Appendix B','Glossary of Terms','37'),
        ('Appendix C','Support and Contacts','39'),
    ]
    for num,title,pg in toc_items:
        p=doc.add_paragraph(); sp(p,0,20)
        lft=Cm(0 if not num.startswith(' ') else 0.5)
        p.paragraph_format.left_indent=lft
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
        run(p,f'{num}  {title}\t{pg}',num in [t[0] for t in [('1',''),('2',''),('3',''),('4',''),('5',''),
            ('6',''),('7',''),('8',''),('9',''),('10',''),('11',''),('12',''),
            ('13',''),('14',''),('Appendix A',''),('Appendix B',''),('Appendix C','')]],
            9 if num.startswith(' ') else 9.5, DB if not num.startswith(' ') else BD)

    # ── CHAPTER 1: INTRODUCTION ───────────────────────────────────────────────
    doc.add_page_break()
    chapter_banner(doc,'CHAPTER 1: INTRODUCTION TO PSFGS')

    section_label(doc,'1.1  About the System')
    para(doc,'PSFGS (Public Sector Financial Governance Suite) is a cloud-based, MFMA-native software platform designed exclusively for South African municipalities. It provides an integrated suite of 13 modules covering the full spectrum of municipal financial governance — from audit finding management to Section 71 reporting, SCM compliance, consequence management, and executive reporting.')
    para(doc,'The platform was developed in response to a persistent challenge: despite significant investment in audit support and training, the number of municipalities achieving clean audit opinions has remained below 30 for six consecutive years. The root cause is structural — municipalities lack the real-time, integrated tooling needed to track findings, implement corrective action, and produce verifiable evidence between audit cycles.')
    para(doc,'PSFGS addresses this gap by digitising every governance workflow that AGSA auditors assess, creating a single source of truth that eliminates the spreadsheet-and-email approach that causes repeat findings.')

    section_label(doc,'1.2  System Architecture')
    for item in [
        ('Cloud-hosted SaaS (Software-as-a-Service)','no local server or hardware installation required'),
        ('Multi-tenant architecture','each municipality\'s data is completely isolated from all others'),
        ('Role-based access control','8 user roles, each with tailored views and permission levels'),
        ('MFMA-native design','every workflow maps to a specific legislative requirement'),
        ('Modern technology stack','FastAPI backend and React frontend — works on any modern browser'),
        ('Responsive design','accessible and fully functional on desktop, tablet, and mobile devices'),
    ]: bullet(doc,item[1],bold_prefix=item[0])

    section_label(doc,'1.3  Supported Browsers and Devices')
    for b in [
        'Google Chrome (recommended, version 100 or later)',
        'Microsoft Edge (version 90 or later)',
        'Mozilla Firefox (version 95 or later)',
        'Safari (version 14 or later — macOS and iOS)',
        'Mobile: iOS Safari and Android Chrome fully supported',
        'Minimum screen resolution: 1280 x 720 pixels',
    ]: bullet(doc,b)

    section_label(doc,'1.4  Logging In — Step by Step')
    for n,t in [
        (1,'Open your web browser and navigate to your municipality\'s PSFGS URL (this is provided by your System Administrator on first setup)'),
        (2,'The login page displays your municipality\'s official logo and name — confirming you are on the correct system'),
        (3,'Enter your email address (e.g. cfo@buffalocity.gov.za)'),
        (4,'Enter your password. The default first-login password is Admin@2026! — you will be prompted to change it on first use'),
        (5,'Click "Sign In"'),
        (6,'You are redirected to the Executive Dashboard — your governance home page'),
    ]: step(doc,n,t)
    callout(doc,'Note:','If you see an error message after clicking Sign In, check that your email address is correct and that Caps Lock is not active. Contact your System Administrator if you cannot log in after two attempts.',bg=LB,border=MB)

    section_label(doc,'1.5  Navigating the System')
    para(doc,'The PSFGS interface has two main areas: the LEFT SIDEBAR (always visible, contains navigation, logo, entity name and logout) and the MAIN CONTENT AREA (changes based on the selected module).',bef=0,aft=40)
    para(doc,'Sidebar Sections:',True,9.5,DB,bef=0,aft=20)
    for section,items in [
        ('Overview','Dashboard (governance home page), Finding Workflow diagram'),
        ('Domain A - Audit Readiness','Audit Findings Tracker (AFT), Asset Verification System (AVS)'),
        ('Domain B - Financial Control','Budget Monitor (BVM), IFWE Register (IFW)'),
        ('Domain C - Reporting','Executive Portal (ERP), Section 71 Auto-Generator (S71), Performance KPIs (PKD)'),
        ('Domain D - Compliance','SCM Compliance (SCC), Consequence Management (CMT), Delegations (DAR), Policy Compliance (PCR)'),
        ('Bottom','Settings (System Admin only), Logout'),
    ]: bullet(doc,items,bold_prefix=section)

    section_label(doc,'1.6  User Roles and Permissions')
    para(doc,'PSFGS uses eight predefined user roles. Each role determines which modules are accessible and what level of access is granted. Roles are assigned by the System Administrator.',bef=0,aft=60)
    roles_data=[
        ('System_Admin','IT Administrator','System config, user management, entity settings','All modules + Settings','Full Admin access'),
        ('Accounting_Officer','City Manager / MM','Strategic oversight, finding assignment, consequence management','AFT, CMT, DAR, Dashboard, ERP','Full write access'),
        ('CFO','Chief Financial Officer','Financial oversight, Section 71 approval, IFWE management','BVM, S71, IFW, Dashboard, ERP','Full write access'),
        ('Director','Directorate Head','Directorate-level implementation, KPI reporting','AFT, AVS, PKD, DAR','Write within directorate'),
        ('Manager','Senior Manager','Operational implementation of corrective actions','AFT, AVS, SCC','Write within area'),
        ('Officer','Operational Staff','Evidence submission, asset verification','AFT (evidence), AVS','Limited write'),
        ('Auditor','Internal Audit Official','Finding capture, review, PCR oversight','AFT (all stages), PCR','Audit-level write'),
        ('Councillor','Elected Councillor','Oversight and governance monitoring','Dashboard, ERP','Read only'),
    ]
    rt=doc.add_table(len(roles_data)+1,5)
    rt.alignment=WD_TABLE_ALIGNMENT.CENTER
    col_w=[Cm(3.5),Cm(3.5),Cm(5),Cm(4),Cm(2.5)]
    hdrs=['Role','Who This Is','Key Responsibilities','Primary Modules','Access Level']
    for i,(h,w) in enumerate(zip(hdrs,col_w)):
        c=rt.rows[0].cells[i]; c.width=w; shade(c,DB); no_border(c); cell_mar(c,60,60,70,70)
        p=c.add_paragraph(); sp(p,0,0); run(p,h,True,8,WH)
    for ri,row in enumerate(roles_data):
        for ci,(v,w) in enumerate(zip(row,col_w)):
            c=rt.rows[ri+1].cells[ci]; c.width=w
            shade(c,LB if ri%2==0 else WH); no_border(c); cell_mar(c,50,50,70,70)
            p=c.add_paragraph(); sp(p,0,0); run(p,v,ci==0,8.5,DB if ci==0 else BD)
    doc.add_paragraph()

    # ── MODULES ────────────────────────────────────────────────────────────────

    module_chapter(doc,2,'DASH','EXECUTIVE DASHBOARD',
        purpose_paras=[
            'The Executive Dashboard is the home screen of PSFGS. It provides a consolidated, real-time overview of the municipality\'s governance health across all 13 modules. Senior management can instantly see the status of audit findings, budget performance, SCM compliance, IFWE totals, and KPI achievement — all without navigating to individual modules.',
            'The Dashboard is designed to be the first thing an Accounting Officer, CFO, or Mayor views each morning. Colour-coded tiles make it immediately clear which governance areas need attention, and the AI analysis panel translates the data into plain-language recommendations.',
        ],
        legislation='MFMA Section 71 (monthly reporting obligations), Section 72 (mid-year assessment), MSA Section 41 (performance management systems).',
        users=[
            ('Municipal Manager / Accounting Officer','Primary user. Reviews overall governance health daily before making decisions or attending meetings.'),
            ('Mayor and Speaker','Weekly governance review for Council reporting preparation and political oversight.'),
            ('CFO','Monitors financial metrics — budget variance, IFWE totals, Section 71 submission status, and overdue findings.'),
            ('Councillors','Read-only access for oversight and accountability purposes.'),
        ],
        features=[
            ('KPI Tiles','Total open findings (colour-coded by severity), Total IFWE amount, SCM compliance rate (%), Active CMT cases, Budget variance alert count, Overdue finding count'),
            ('AI Governance Health Analysis','Automated narrative assessment identifying the top 3 governance risks and recommended actions — updated on each login'),
            ('Module Navigation Shortcuts','Click any KPI tile to navigate directly to the relevant module for detailed investigation'),
            ('Colour-coded Status Indicators','Green = within acceptable range; Amber = attention required; Red = critical action needed today'),
            ('Real-time Data','Dashboard refreshes automatically on every login and every 30 minutes during an active session'),
            ('Entity Branding','Municipality logo, name, and abbreviation displayed in the sidebar throughout the session'),
        ],
        steps=[
            (1,'Log in to PSFGS — the Executive Dashboard loads automatically as your home page'),
            (2,'Review the top row of KPI tiles — note any tiles showing red or amber status before doing anything else'),
            (3,'Read the AI Governance Health Analysis panel — this provides a written summary of the current top risks in plain language'),
            (4,'Click on any KPI tile to navigate directly to the relevant module for detailed investigation and corrective action'),
            (5,'Return to the Dashboard at any time by clicking "Dashboard" in the left sidebar'),
            (6,'Use the Dashboard before every Council meeting, MPAC briefing, or management committee sitting to get a current, accurate picture of the municipality\'s governance status'),
        ],
        tip_label='Best Practice:',
        tip_text='The Dashboard is your governance vital signs monitor. Make it a daily habit to check it before starting work. A red tile means action is required today — do not navigate away without understanding the cause. The AI analysis panel is especially useful before Council meetings as it summarises governance risks in plain language that can be presented without preparation.',
        errors=[
            ('"Dashboard tiles show zero"','No data has been loaded for your entity in the database','Contact the System Administrator to verify that entity_id is correctly configured and that seed data has been loaded'),
            ('"AI Analysis not loading"','API timeout or slow internet connection','Refresh the page; if the problem persists check your internet connection and try again after 2 minutes'),
            ('"Logo not showing in sidebar"','Logo has not been uploaded in Entity Settings','System Administrator must upload the municipal logo via Settings > Entity Information > Logo Upload'),
        ]
    )

    module_chapter(doc,3,'AFT','AUDIT FINDINGS TRACKER',
        purpose_paras=[
            'The Audit Findings Tracker (AFT) is the core module of PSFGS. It manages the complete lifecycle of every audit finding — from the moment it is captured by Internal Audit to the day it is formally closed with verified evidence. The AFT supports both AGSA external audit findings and internal audit findings, ensuring corrective actions are tracked, assigned, implemented, evidenced, and reviewed in a structured, auditable process.',
            'The AFT eliminates the most common cause of repeat findings: the loss of tracking between audit cycles. Every finding has a named owner, a due date, a status, and an evidence trail — all visible to every authorised user in real time.',
        ],
        legislation='MFMA Section 131 (implementation of audit recommendations), Public Audit Act 25 of 2004, AGSA Findings Management Protocol, MFMA Section 62(1)(c) (internal controls).',
        users=[
            ('Auditor / Internal Audit','Captures new findings from AGSA management letters or internal audit reports; reviews submitted evidence; updates status to Under_Review'),
            ('Accounting_Officer / City Manager','Assigns findings to responsible officials; sets resolution due dates; adds directive notes; signs off on high-severity closures'),
            ('Director / Manager','Receives assigned findings; implements corrective action; updates status to In_Progress'),
            ('Officer','Submits documentary evidence once corrective action is complete; updates status to Evidence_Submitted'),
            ('CFO','Reviews financial-impact findings; approves closure of Critical and High severity findings'),
        ],
        features=[
            ('Finding Capture','AG Reference Number, description, finding type (Compliance/Financial/Performance/Operational), severity (Critical/High/Medium/Low), financial impact amount, directorate, and audit year'),
            ('6-Stage Lifecycle','New > Assigned > In_Progress > Evidence_Submitted > Under_Review > Closed — plus automatic Overdue escalation after 90 days'),
            ('Assignment Modal','Assign to any system user with a dropdown showing name, role and directorate; set resolution due date; add written directive notes'),
            ('Assigned To Column','Every finding row shows the current assignee name and due date — or "Unassigned" if not yet assigned'),
            ('Repeat Finding Flag','Mark findings as repeats from prior years; track the number of consecutive years the finding has appeared'),
            ('Statistics Tiles','Total findings, Open count, Closed count, Overdue count, Critical severity count — live on the findings page'),
            ('Filters','Filter the findings list by audit year, status, severity, or directorate for targeted management'),
        ],
        steps=[
            (1,'CAPTURE — Log in as Auditor > Click "Audit Findings (AFT)" in the sidebar > Click "+ New Finding"'),
            (2,'FILL FORM — Enter AG Reference Number exactly as it appears in the AGSA management letter > Select the Audit Year > Write a full, clear description of the finding > Select Finding Type, Severity, Financial Impact, and Directorate > Tick "Repeat Finding" if the same issue was raised in a prior year > Click "Save Finding" — status is set to New'),
            (3,'ASSIGN — Log in as Accounting_Officer > Find the New finding > Click the "Assign" button > Select the responsible official from the dropdown (shows name, role, directorate) > Set Resolution Due Date > Type directive notes (what corrective action is required) > Click "Assign Finding" — status automatically changes to Assigned'),
            (4,'IMPLEMENT — Log in as Director or Manager > Find your assigned finding > Change the status dropdown to "In_Progress" > Begin implementing the corrective action as directed'),
            (5,'SUBMIT EVIDENCE — Once corrective action is complete, change status to "Evidence_Submitted" > Reference the evidence documents in the notes field (e.g. file name, date, Council resolution number)'),
            (6,'REVIEW — Log in as Auditor or IA Director > Review findings in Evidence_Submitted status > Assess whether evidence adequately demonstrates that the root cause has been addressed > Change status to "Under_Review" if evidence is sufficient'),
            (7,'CLOSE — Log in as CFO or MM > Review the finding, evidence, and audit trail > If fully satisfied, change status to "Closed" > Finding is removed from the open count on the Dashboard'),
        ],
        tip_label='Critical Rule:',
        tip_text='Always capture the AG Reference Number exactly as it appears in the AGSA management letter — this is the primary identifier used for tracking across financial years and for evidence submission to AGSA. Never close a finding without adequate supporting evidence; premature closure is a high-risk audit exposure. Findings that reach 90 days without a status update automatically escalate to Overdue and trigger notifications to the MM and CFO.',
        errors=[
            ('"Error creating finding" on save','Audit year not selected (set to zero) or description is empty','Select a valid audit year from the dropdown before saving; ensure the description field is not blank'),
            ('"Assigned To" shows Unassigned','The Assign modal was not used — status changed directly without assigning','Click the "Assign" button on the finding row and complete the assignment form'),
            ('"Cannot change status" for a user','The user\'s role does not have permission for that lifecycle stage transition','Use the correct user account: each stage requires a specific role (see Section 1.6)'),
            ('"Finding not visible" to a Director','Finding not linked to that Director\'s directorate','MM must use the Assign modal to re-assign the finding to the correct directorate official'),
        ]
    )

    module_chapter(doc,4,'AVS','ASSET VERIFICATION SYSTEM',
        purpose_paras=[
            'The Asset Verification System enables the systematic physical verification of municipal assets against the fixed asset register. It supports compliance with GRAP 17 requirements for property, plant and equipment, and provides a structured, evidence-based process for quarterly and annual verification exercises.',
            'Unverified or missing assets are among the most common AGSA findings, often resulting in material misstatements in the Annual Financial Statements. The AVS ensures that every asset is physically located, its condition assessed, and any discrepancies reported and investigated.',
        ],
        legislation='MFMA Section 63 (asset and liability management), GRAP 17 (Property, Plant and Equipment), Municipal Asset Transfer Regulations 2008, National Treasury Asset Management Guidelines.',
        users=[
            ('Asset Manager','Creates verification sessions; assigns verification teams; reconciles results against the asset register; reports discrepancies'),
            ('Finance Officer / Manager','Conducts physical verification in the field; captures condition scores and GPS coordinates'),
            ('CFO','Reviews discrepancy reports; approves write-offs and disposals; refers missing assets to Internal Audit'),
            ('Director','Oversees directorate asset verification; approves condition assessments for their functional area'),
            ('Auditor','Uses verification session results as audit evidence for GRAP 17 compliance testing'),
        ],
        features=[
            ('Asset Register','Complete list of municipal assets with descriptions, acquisition values, carrying values, and physical locations'),
            ('Verification Sessions','Create sessions filtered by directorate, physical location, or asset category'),
            ('Condition Scoring','Good / Fair / Poor / Missing / Stolen — mandatory notes required for Poor condition and below'),
            ('GPS Tagging','Record the physical GPS coordinates of each asset at time of verification for future reference'),
            ('Discrepancy Reports','Automatic identification of unverified assets, condition changes from prior verification, and missing or stolen assets'),
            ('GRAP 17 Alignment','Condition scores feed into useful life assessments and depreciation calculations for AFS preparation'),
            ('Verification Statistics','Percentage of assets verified, condition distribution by category, and total asset value by condition grade'),
        ],
        steps=[
            (1,'Log in as Asset Manager or Director > Click "Asset Verification (AVS)" in the sidebar'),
            (2,'Click "New Verification Session" > Enter a descriptive session name (e.g. "Q2 2025/26 EIS Vehicle Verification") > Select the directorate > Set the verification start and end dates'),
            (3,'Assign the verification team members who will conduct the physical verification'),
            (4,'Team members physically locate each asset in the field using the asset register list as a checklist'),
            (5,'For each located asset: enter the condition score (Good/Fair/Poor) > record GPS coordinates > add any observations or notes'),
            (6,'Mark each verified asset as "Verified" in the system — unverified assets remain highlighted for follow-up'),
            (7,'Once all assets are verified or confirmed missing, close the session — the system generates an automatic discrepancy report'),
            (8,'CFO reviews the discrepancy report > assets marked Missing are immediately referred to Internal Audit for investigation'),
            (9,'Generate the GRAP 17 compliance verification report for use in AFS preparation and AGSA submission'),
        ],
        tip_label='Compliance Requirement:',
        tip_text='Asset verification must be conducted at least quarterly per National Treasury guidelines — AGSA specifically tests whether the asset register reflects physical reality. Any asset marked Missing or Stolen must be reported to the MM and SAPS within 24 hours of discovery, and captured in the IFWE register if the loss has a financial value. Do not close a verification session without accounting for every asset on the list.',
        errors=[
            ('"Asset not on list"','The asset has not been captured in the asset register','Add the asset to the asset register first with all required details, then include it in the verification session'),
            ('"Session shows Incomplete"','Some assets in the session have not yet been verified','Complete the physical verification of all outstanding assets before attempting to close the session'),
            ('"GPS not capturing coordinates"','Browser location permissions are blocked','Enable location access for PSFGS in your browser\'s security settings and reload the page'),
        ]
    )

    module_chapter(doc,5,'BVM','BUDGET VS ACTUAL MONITOR',
        purpose_paras=[
            'The Budget vs Actual Monitor provides real-time visibility of expenditure against all approved budget votes. It enables the CFO, directors, and the Municipal Manager to identify variances early, understand their causes, and take corrective action before variances compound, trigger Section 71 disclosures, or become AGSA findings.',
            'The BVM is also the data source for the Section 71 Auto-Generator — ensuring that monthly National Treasury reports are compiled from the same live data used for management monitoring.',
        ],
        legislation='MFMA Section 71 (monthly expenditure reporting), Section 72 (mid-year budget and performance assessment), Municipal Budget and Reporting Regulations 2009, NT Circular 88 on budget monitoring.',
        users=[
            ('CFO','Primary user — monitors all budget votes; flags variances; approves budget adjustments and virements'),
            ('Directors','Monitor their own directorate votes; explain variances; submit virement requests'),
            ('Finance Officers','Capture actual expenditure each month; reconcile with financial system records'),
            ('Municipal Manager','Strategic oversight of overall budget performance; escalation point for material variances'),
            ('Council / MPAC','Quarterly review of budget performance reports; oversight of underspending on capital projects'),
        ],
        features=[
            ('Budget Vote Dashboard','All votes displayed with approved budget, actual expenditure to date, variance in Rand, and variance as a percentage'),
            ('Variance Alerts','Automatic red flag for overspend above 10%; amber warning for underspend above 25% (critical at mid-year for capital)'),
            ('Directorate Filter','View all votes for a specific directorate; drill from municipal total to individual vote level'),
            ('Period Filter','Select any month or cumulative quarter within the financial year for comparative analysis'),
            ('Trend Charts','12-month expenditure trend per vote — helps distinguish timing differences from genuine variances'),
            ('Section 71 Data Feed','BVM data flows directly and automatically into the Section 71 module — no re-entry required'),
            ('Virement Tracking','Record approved budget virements with reference to the authorising Council resolution or delegation'),
        ],
        steps=[
            (1,'Log in as CFO or Director > Click "Budget Monitor (BVM)" in the sidebar'),
            (2,'Review the overview panel — total approved budget, total actual expenditure, overall variance percentage for the current period'),
            (3,'Use the Directorate filter to drill into a specific directorate\'s budget performance'),
            (4,'Use the Period filter to select the current month or cumulative quarter'),
            (5,'Identify votes flagged red (overspend) or amber (approaching the variance threshold)'),
            (6,'Click on any flagged vote to see the detailed expenditure breakdown by cost element'),
            (7,'Request a written variance explanation from the responsible Director for any variance above 10%'),
            (8,'Record the explanation and the agreed corrective action in the vote\'s notes field'),
            (9,'If a virement is required to correct an allocation: capture the virement request, reference the approving resolution, and update both affected votes'),
        ],
        tip_label='MFMA Compliance:',
        tip_text='Variances above 10% require a written explanation to the Municipal Manager per MFMA Regulation 28 — this requirement applies every month, not just at year-end. Do not allow variances to accumulate month on month; address them immediately when first identified. Underspending above 25% on capital votes at mid-year (December) triggers a mandatory disclosure in the Section 72 mid-year assessment. Use the BVM trend chart to distinguish genuine overspend from legitimate timing differences such as annual insurance premiums paid in a single month.',
        errors=[
            ('"All votes showing zero actual"','Actual expenditure has not been captured for the current period','Finance staff must capture actuals from the financial system for the current period before using BVM'),
            ('"Variance % showing as incorrect"','Budget amount for that vote is zero or blank','Capture the full approved budget for all votes at the start of the financial year before monitoring begins'),
            ('"S71 not pulling BVM data"','Period selected in S71 does not match the period with data in BVM','Ensure the same financial period is selected in both BVM and the Section 71 module'),
        ]
    )

    module_chapter(doc,6,'IFW','IFWE REGISTER',
        purpose_paras=[
            'The IFWE Register captures, classifies, tracks, and manages all Irregular, Fruitless and Wasteful Expenditure identified within the municipality. It provides a structured workflow from initial identification through investigation, determination of responsibility, application for condonation or recovery, and final disclosure in the Annual Financial Statements.',
            'Maintaining a complete and current IFWE register is both a legal obligation under the MFMA and a key determinant of audit outcomes. Incomplete registers, or registers that do not reconcile with prior year AFS disclosures, are consistently raised as findings by the AGSA.',
        ],
        legislation='MFMA Section 32 (irregular expenditure), Section 34 (fruitless and wasteful expenditure), Section 38(1)(c)(ii) (prevention of irregular expenditure), Treasury Regulation 9.1, National Treasury IFWE Framework 2019.',
        users=[
            ('CFO','Primary register owner — responsible for completeness; approves condonation applications; presents to Council'),
            ('Internal Audit','Identifies IFWE during audits and reviews; refers items to CFO for capture; investigates circumstances'),
            ('Finance Director','Captures individual expenditure items; tracks investigation progress; maintains historical records'),
            ('Municipal Manager','Approves write-offs; institutes consequence management for responsible officials; refers to Council'),
            ('Council / MPAC','Oversight body — receives the Section 32(2) report; approves or rejects condonation applications'),
        ],
        features=[
            ('IFWE Classification','Irregular (non-compliance with procurement or process requirements) / Fruitless (no value received) / Wasteful (avoidable expenditure)'),
            ('Amount Capture','Rand value, financial year incurred, vote number, supplier or payee, and circumstances description'),
            ('Investigation Status','Identified > Under Investigation > Awaiting Council Decision > Condoned / Recovered / Written Off'),
            ('Condonation Tracking','Record application submitted to NT or Council with date; outcome and approval reference recorded when received'),
            ('AFS Disclosure Schedule','Running totals by financial year and classification — ready for direct use in AFS notes preparation'),
            ('Council Report Generator','Auto-generates the Section 32(2) IFWE disclosure report for submission to Council within the 10-day deadline'),
            ('Historical Register','All prior year IFWE maintained permanently for AGSA comparison and year-on-year reconciliation'),
        ],
        steps=[
            (1,'IDENTIFY — Any official who identifies a transaction that may constitute IFWE must report it to the CFO immediately, verbally and in writing'),
            (2,'CAPTURE — CFO or Finance Director logs in > Click "IFWE Register (IFW)" > Click "+ New Item"'),
            (3,'CLASSIFY — Select the expenditure type (Irregular, Fruitless, or Wasteful) > Enter the Rand amount > Select the financial year and vote number > Write a clear description of the circumstances and why the expenditure qualifies as IFWE'),
            (4,'INVESTIGATE — Assign the investigating official > Set the investigation completion deadline > Status: Under Investigation'),
            (5,'DETERMINE — Record the investigation outcome: identify the responsible official(s); confirm whether the expenditure is recoverable from that official'),
            (6,'COUNCIL REPORT — Generate the Section 32(2) disclosure report > Submit to Council at the next available meeting > Status: Awaiting Council Decision'),
            (7,'RESOLVE — Record the Council decision: Condone (NT approval required if above R500,000) or Recover from responsible official > Update the item status accordingly'),
            (8,'DISCLOSE — Item automatically flows into the AFS disclosure schedule with correct year and classification'),
        ],
        tip_label='Legal Requirement:',
        tip_text='IFWE must be reported to Council within 10 days of discovery per MFMA Section 32(2) — failure to report timeously is itself a separate AGSA finding. Irregular expenditure does NOT automatically mean money was misused; it means the correct procurement process was not followed. However, every rand must be investigated and either condoned with NT approval or recovered from the responsible official. Never delete an IFWE item from the register — AGSA compares the register to every prior year AFS disclosure and will raise any unexplained removal.',
        errors=[
            ('"Running total does not match AFS"','Prior year IFWE items not carried forward into the current year register','Ensure all historical IFWE from prior AFS disclosures is captured in the register including opening balances'),
            ('"Condonation status not updating"','NT or Council decision received but not captured in system','Update the resolution status and reference number immediately upon receiving the official decision'),
            ('"Item classified incorrectly"','Misclassification between Irregular and Fruitless/Wasteful','Edit the item, correct the classification, and add a note explaining the amendment for the audit trail'),
        ]
    )

    module_chapter(doc,7,'SCC','SCM COMPLIANCE CHECKER',
        purpose_paras=[
            'The SCM Compliance Checker automates the compliance testing of procurement requisitions against the municipality\'s Supply Chain Management policy and applicable legislation. It significantly reduces the risk of irregular expenditure by identifying non-compliance before a procurement commitment is made — not after the AGSA raises it as a finding.',
            'The automated rule engine checks 12 key compliance points on every requisition, scores the result, and produces a prioritised exceptions list with remediation guidance for each failed check.',
        ],
        legislation='MFMA SCM Regulations 2005, Preferential Procurement Policy Framework Act (PPPFA) 5 of 2000, National Treasury Instruction Notes 3 and 7 of 2021/22, MFMA Section 112 (SCM policy).',
        users=[
            ('SCM Director','Primary user — reviews compliance scores; approves requisitions above the threshold; manages exception escalations'),
            ('Demand Manager','Captures requisitions in the system; runs compliance checks; resolves identified exceptions before approval'),
            ('CFO','Reviews the exceptions summary report; approves high-value or non-compliant requisitions by exception'),
            ('Internal Audit','Monitors SCM compliance trends on the Dashboard; uses the exceptions report as audit evidence during reviews'),
        ],
        features=[
            ('Requisition Capture','Description, estimated rand value, supplier name, motivation, and SCM procurement method (quotation/competitive bid/sole source)'),
            ('12-Point Rule Engine','Automated checks: threshold classification, number of quotes obtained, tax clearance validity, CSD registration, BEE compliance, sole-source justification, conflict of interest declaration, specification adequacy'),
            ('Compliance Score','0-100% score with colour coding: Green (80-100% = proceed), Amber (60-79% = resolve before approval), Red (below 60% = CFO review required)'),
            ('Exceptions List','Itemised list of failed checks with the specific MFMA regulation reference for each exception'),
            ('Remediation Guidance','Each exception includes clear, actionable guidance on how to resolve it before resubmission'),
            ('Trend Reporting','Monthly SCM compliance rate for the municipality flows to the Dashboard and is tracked over time'),
        ],
        steps=[
            (1,'Log in as SCM Director or Demand Manager > Click "SCM Compliance (SCC)" in the sidebar'),
            (2,'Click "+ New Requisition" > Enter all requisition details: description, estimated value, supplier, motivation, and procurement method'),
            (3,'Select the SCM method (quotation / competitive bid / sole source) and complete all method-specific fields'),
            (4,'Click "Run Compliance Check" — the system tests all 12 rules automatically in seconds'),
            (5,'Review the compliance score and the exceptions list — read each exception carefully including the regulation reference'),
            (6,'For each failed check: follow the remediation guidance to resolve the underlying compliance issue'),
            (7,'Once all exceptions are resolved, click "Run Compliance Check" again to confirm the score has improved to above 80%'),
            (8,'SCM Director reviews the final score and approved requisition > requisition proceeds to purchase order processing'),
            (9,'Any requisition scoring below 60% must be reviewed and approved by the CFO before any further processing'),
        ],
        tip_label='Prevention First:',
        tip_text='Never approve a requisition that has active exceptions — doing so is the direct cause of irregular expenditure in most municipalities. The compliance checker catches the most common procedural errors but officials remain personally liable for compliance under the MFMA. All sole-source motivations must be approved in writing by the CFO before the requisition is captured in the system.',
        errors=[
            ('"All 12 rules failing on check"','Requisition is missing required information in key fields','Complete all mandatory fields in the requisition form before running the compliance check'),
            ('"Tax clearance check failing for a registered supplier"','CSD record is outdated or tax compliance pin has expired','Verify the supplier\'s CSD registration is active and that the tax compliance pin is valid on SARS eFiling'),
            ('"Score drops lower on re-check"','Resolving one exception exposed an underlying compliance gap','Read all exceptions carefully; address the root cause of each issue, not just the surface symptom'),
        ]
    )

    module_chapter(doc,8,'CMT','CONSEQUENCE MANAGEMENT TRACKER',
        purpose_paras=[
            'The Consequence Management Tracker provides a structured digital system for managing all financial misconduct cases, disciplinary proceedings, AGSA referrals, and criminal matters arising from MFMA contraventions within the municipality. It ensures that consequence management is visible, tracked, progressed, and reported to Council.',
            'The absence of effective consequence management is one of the AGSA\'s most persistent findings. PSFGS ensures that every case is opened formally, assigned to an investigating officer, tracked to completion, and reported — eliminating the "no outcome" situation that results in adverse audit opinions.',
        ],
        legislation='MFMA Section 171-173 (financial misconduct proceedings), Municipal Systems Act Section 67 (disciplinary procedures), COGTA Consequence Management Framework 2020, National Treasury Consequence Management Directive.',
        users=[
            ('Municipal Manager','Opens and monitors cases; institutes formal proceedings against directors and senior managers; reports to Council'),
            ('HR Director','Manages disciplinary processes; schedules hearings; captures outcomes and sanctions imposed'),
            ('Legal Advisor','Tracks criminal referrals to SAPS or NPA; manages civil recovery proceedings'),
            ('Internal Audit','Identifies financial misconduct and refers to the MM; tracks AGSA-referred cases and required outcomes'),
            ('CFO','Reviews cases with financial recovery implications; approves civil recovery actions'),
        ],
        features=[
            ('Case Capture','Respondent details, case type, MFMA contravention reference, description of alleged misconduct, date of discovery'),
            ('Case Classification','Financial Misconduct, Gross Negligence, Fraud, Theft, Corruption, Disciplinary — each with specific legislative requirements'),
            ('Investigating Officer','Assign investigating officer; set investigation deadline; track progress updates'),
            ('Hearing Management','Schedule hearing dates; capture chairperson; record union representation; capture outcome and sanction imposed'),
            ('AGSA Referral Tracking','Track cases referred by AGSA for follow-up; outcomes required within 60 days of AGSA notification'),
            ('Appeal Management','Record appeal lodgement date, appeal body, and final appeal outcome'),
            ('Council Reporting','Auto-generates the quarterly consequence management report for Council and MPAC'),
        ],
        steps=[
            (1,'IDENTIFY — Misconduct is identified by Internal Audit, an AGSA finding, a whistle-blower, or management observation'),
            (2,'OPEN CASE — Log in as MM or HR Director > Click "Consequence Management (CMT)" > Click "+ New Case"'),
            (3,'CAPTURE DETAILS — Enter respondent name, role, and directorate > Select the case type > Write a clear, factual description of the alleged contravention > Reference the specific MFMA section breached'),
            (4,'ASSIGN — Appoint an Investigating Officer > Set the investigation completion deadline'),
            (5,'INVESTIGATE — Investigating officer conducts the investigation; preliminary findings are reported to the MM'),
            (6,'SCHEDULE HEARING — Schedule the disciplinary hearing > Capture the date, chairperson, and whether union representation has been arranged'),
            (7,'RECORD OUTCOME — After the hearing: capture outcome (Guilty / Not Guilty / Withdrawn) > Record sanction imposed (warning, demotion, dismissal, criminal referral, civil recovery)'),
            (8,'APPEAL — If the respondent appeals: capture the appeal date and the appeal body (GPSSBC, Labour Court, etc.) > Record the final appeal outcome'),
            (9,'CLOSE — Record the final case status > Case feeds into the quarterly Council report automatically'),
            (10,'REPORT — Generate and submit the quarterly consequence management report to Council/MPAC'),
        ],
        tip_label='Critical Timelines:',
        tip_text='Cases older than 180 days without a hearing outcome must be reported to COGTA under the Consequence Management Framework — this is monitored nationally. AGSA specifically tests whether cases from prior year findings have been finalised. A municipality with stale, open cases will receive a finding on inadequate consequence management. Never close a case without capturing the actual outcome — "under investigation" is a status, not an outcome.',
        errors=[
            ('"Case not in Council report"','Case status not set to Active or Open','Update the case status to Active and ensure the quarter is set correctly'),
            ('"AGSA referral count wrong on Dashboard"','Referral source not marked as AGSA on the case','Edit the case record and set the referral source field to AGSA'),
            ('"Hearing date passed — no outcome recorded"','Staff did not capture the hearing outcome after the hearing was held','Update the case record with the outcome immediately after every hearing — do not leave it blank'),
        ]
    )

    module_chapter(doc,9,'DAR','DELEGATION OF AUTHORITY REGISTER',
        purpose_paras=[
            'The Delegation of Authority Register maintains a complete, current, and auditable digital register of all powers delegated by Council to the Municipal Manager, and all sub-delegations granted by the MM to senior officials. It ensures that every financial approval decision within the municipality is backed by a valid, in-force delegation with a clear rand limit and an enabling Council resolution.',
            'Operating without valid delegations — or approving transactions beyond the limits of existing delegations — constitutes irregular expenditure and is a direct AGSA finding. The DAR module prevents this by making delegation status immediately visible and alerting officials when delegations are approaching expiry.',
        ],
        legislation='MFMA Section 79 (delegations by municipal councils), MSA Section 59 (sub-delegations by the MM), MFMA Section 82 (limitation on delegations), National Treasury Delegation of Authority Guidelines for Local Government.',
        users=[
            ('Municipal Manager','Primary owner — receives Council delegations; grants sub-delegations to directors and the CFO'),
            ('CFO','Receives financial management delegations; manages financial approval limits for the finance directorate'),
            ('Directors','Receive operational delegations for their functional areas; grant further sub-delegations to managers'),
            ('Company Secretary','Maintains the register; ensures all Council resolutions are correctly referenced; schedules annual reviews'),
            ('Internal Audit','Reviews delegation coverage for gaps; tests during the year that approvals are within valid delegation limits'),
        ],
        features=[
            ('Delegation Hierarchy','Full Council > MM > CFO/Directors > Managers chain visible in a single view'),
            ('Financial Limits','Rand value approval limit clearly stated for every financial delegation'),
            ('Enabling Resolution','Reference to the specific Council resolution number and date authorising each delegation'),
            ('Validity Period','Start and end dates; automatic expiry alert 30 days before any delegation expires'),
            ('Coverage Gap Analysis','Identifies municipal functions that do not have an active delegation in place'),
            ('PDF Export','Generate a formatted delegation letter for signing and filing'),
            ('Annual Review','Flags all delegations for review at the start of each new financial year'),
        ],
        steps=[
            (1,'COUNCIL RESOLUTION — Council passes a resolution delegating powers to the MM, or the MM sub-delegates to officials'),
            (2,'CAPTURE — Log in as MM or Company Secretary > Click "Delegations (DAR)" > Click "+ New Delegation"'),
            (3,'SELECT PARTIES — Choose the Delegator (e.g. Council or MM) and the Delegatee (e.g. CFO or Director)'),
            (4,'DESCRIBE POWER — Write a clear description of the power being delegated > Reference the enabling Council resolution number and date'),
            (5,'SET LIMITS — Enter the Financial Limit (Rand value) > Set Start Date and Expiry Date (or mark as ongoing for non-financial delegations)'),
            (6,'SAVE — System saves the delegation and sends a notification to the delegatee'),
            (7,'ANNUAL REVIEW — At the start of each financial year: review all delegations > update financial limits if changed > renew expiry dates > obtain a new Council resolution if the power itself is changing'),
            (8,'ONGOING MONITORING — Internal Audit tests during the year that all approved transactions are within the limits of valid delegations'),
        ],
        tip_label='Ultra Vires Risk:',
        tip_text='Any decision made without a valid, in-force delegation is ultra vires and constitutes irregular expenditure under the MFMA. Delegations must be renewed at the start of every financial year. If the delegator changes (e.g. a new MM is appointed), all delegations granted by the previous MM must be re-issued. AGSA tests delegation completeness, currency, and compliance at every audit — an incomplete register is a guaranteed finding.',
        errors=[
            ('"Delegation showing as Expired"','The end date has passed without the delegation being renewed','Update the end date, obtain a new Council resolution if required, and save the updated delegation'),
            ('"Financial limit field is blank"','Delegation was captured without entering a rand value','Edit the delegation record and add the specific rand limit — no financial delegation should exist without one'),
            ('"Delegatee not notified"','Email address is not correctly set in the user\'s profile','System Administrator must verify and update the official\'s email address in the user management settings'),
        ]
    )

    module_chapter(doc,10,'PCR','POLICY COMPLIANCE MANAGER',
        purpose_paras=[
            'The Policy Compliance Manager maintains a comprehensive register of all municipal policies, tracks their review dates, and provides an automatic gap analysis between the policies currently in place and those required by the MFMA and related legislation. Outdated, unapproved, or missing policies are among the most frequently cited AGSA findings across all municipality categories.',
            'The module ensures that no policy falls through the cracks between Council approvals and that every required policy is maintained, reviewed on schedule, and supported by a current Council resolution.',
        ],
        legislation='MFMA Section 62(1)(a) (sound financial management systems), COGTA Policy Framework for Local Government, National Treasury Guidelines on Municipal Financial Management Policies, Municipal Systems Act.',
        users=[
            ('CFO','Owns all financial management policies — SCM Policy, Budget Policy, IFWE Policy, Investment Policy, Tariff Policy'),
            ('HR Director','Owns HR policies, performance management policy, and leave management policy'),
            ('Company Secretary','Maintains the complete policy register; schedules Council approval and review processes'),
            ('Internal Audit','Reviews policy coverage for gaps; raises missing or expired policies as audit findings'),
            ('All Directors','Each responsible for policies within their functional area — confirm review dates and upload approved versions'),
        ],
        features=[
            ('Policy Register','All municipal policies with category (Financial / SCM / HR / IT / Governance / Risk / Performance / Other)'),
            ('Review Schedule','Automatic 30-day alert before each policy\'s review due date — sent to the policy owner'),
            ('Gap Analysis','System lists all MFMA-required policies and flags those that are missing, in draft, or expired'),
            ('Owner Assignment','Each policy linked to a named responsible official and their directorate'),
            ('Council Approval Tracking','Record the Council approval date and resolution reference for each policy'),
            ('Status Tracking','Draft / Active / Under Review / Expired / Suspended — updated throughout the review cycle'),
            ('Document Upload','Attach the approved policy PDF to each register entry for instant access by any user'),
        ],
        steps=[
            (1,'Log in as CFO or Company Secretary > Click "Policy Compliance (PCR)" in the sidebar'),
            (2,'Run the gap analysis — review the list of required policies that are missing, in draft, or showing as expired'),
            (3,'For each existing policy: verify it has a named Owner, a current Review Date, and the correct Status'),
            (4,'When a review-due alert is received (30 days before expiry): assign the review task to the policy owner'),
            (5,'Policy owner reviews and revises the policy document — consults with relevant stakeholders'),
            (6,'Revised policy goes to the relevant management committee for comment, then to Council for formal approval'),
            (7,'Once Council approves: update the Status to Active > upload the approved PDF > record the Council resolution reference and approval date'),
            (8,'Set the next review date > the system will automatically alert 30 days before that date'),
        ],
        tip_label='AGSA Priority Policies:',
        tip_text='AGSA specifically and routinely checks for these policies in every audit: SCM Policy, Budget Policy, Asset Management Policy, Investment Policy, Tariff Policy, and Performance Management Policy. A policy not reviewed in over 3 years is treated as non-compliant even if it technically has an Active status. Every policy must have a named individual as owner — "CFO\'s Office" or "Finance Department" is not an acceptable owner.',
        errors=[
            ('"Policy showing Expired but was recently approved"','Review date and approval date not updated after the Council approval','Update both the review date and the Council approval date immediately after every Council approval'),
            ('"Gap analysis shows a policy as missing"','The policy exists but has not been captured in the PCR register','Locate the existing policy, capture it in the register with all details, and upload the approved document'),
            ('"30-day review alert not received"','The policy owner\'s email address is not set up in their user profile','System Administrator must verify that all policy owners have valid email addresses in their user settings'),
        ]
    )

    module_chapter(doc,11,'S71','SECTION 71 AUTO-GENERATOR',
        purpose_paras=[
            'The Section 71 Auto-Generator automatically compiles monthly financial performance reports in the National Treasury prescribed format, drawing live data directly from the Budget vs Actual Monitor. This module eliminates the high-risk, time-pressured manual compilation process that causes errors and late submissions.',
            'Section 71 reports are a legal requirement under the MFMA. Late or inaccurate submissions are recorded by National Treasury and raised as findings by AGSA. The auto-generator ensures that CFOs spend their time reviewing and approving accurate reports, not manually compiling them.',
        ],
        legislation='MFMA Section 71 (monthly budget statements by accounting officers), Municipal Budget and Reporting Regulations 2009 (Regulation 10), National Treasury Section 71 Report Formats.',
        users=[
            ('CFO','Primary user — reviews the auto-generated report; makes corrections in BVM if needed; formally approves before submission'),
            ('Finance Director','Verifies data accuracy; reconciles with financial system records; escalates discrepancies to CFO'),
            ('Municipal Manager','Reviews and co-signs the Section 71 report where the municipality\'s policy requires MM sign-off'),
            ('Financial Reporting Officer','Manages the submission to National Treasury; tracks submission confirmation and acknowledgements'),
        ],
        features=[
            ('Auto-Compilation','All budget and actual expenditure data pulled directly from BVM — no manual re-entry, no transcription errors'),
            ('NT Prescribed Format','Report generated in the exact format stipulated by National Treasury — submission-ready'),
            ('Prior Period Comparison','Current month compared to prior month and to the same month in the prior financial year'),
            ('CFO Review Workflow','CFO formally reviews, annotates corrections, and approves with a timestamp — creating an audit trail'),
            ('Submission Tracking','Captures the submission date and NT acknowledgement number for every monthly report'),
            ('Late Submission Alert','Automatic warning if the 10th-of-month deadline is within 3 days and no report has been approved'),
            ('PDF and Excel Export','Both formats available for submission to NT and for Council/MPAC distribution'),
        ],
        steps=[
            (1,'Month-end period (1st to 8th of each month) — log in as CFO > Click "Section 71 (S71)" in the sidebar'),
            (2,'Select the reporting period from the period dropdown (e.g. March 2026)'),
            (3,'The system automatically compiles the full Section 71 report using BVM data — review every line item for reasonableness'),
            (4,'Check that all totals, subtotals, and percentages are consistent with your knowledge of the municipality\'s financial performance'),
            (5,'If corrections are needed: navigate to BVM > make the corrections > return to S71 and refresh the report'),
            (6,'Once the report is accurate and complete: click "Approve Report" — the system records the CFO approval with a date and time stamp'),
            (7,'Export the report to PDF > submit to National Treasury via the NT online portal or designated email address by the 10th of the month'),
            (8,'Capture the NT submission confirmation number and submission date in the system for the audit trail'),
            (9,'The approved report is archived and immediately available for AGSA inspection'),
        ],
        tip_label='Submission Deadline:',
        tip_text='The Section 71 submission deadline is strictly the 10th of each month with no extensions. Late submissions trigger a formal Section 71(3) notice from National Treasury and are recorded as a finding at the next AGSA audit. Set a recurring calendar reminder for the 5th of each month to begin the review process — this provides 5 working days to identify and correct discrepancies before the deadline. Never submit a report with unexplained variance discrepancies; NT will query them.',
        errors=[
            ('"Report shows blank lines or zeros"','Some budget votes have no actual expenditure captured in BVM for the period','Finance must ensure all actual expenditure is loaded into BVM before the Section 71 review date'),
            ('"Totals not matching between sections"','Corrections made in BVM have not been saved before viewing S71','Save all BVM changes first, then return to S71 and refresh the report view'),
            ('"PDF export is empty or corrupted"','Browser is blocking file downloads','Allow PSFGS to download files in your browser\'s security settings; Google Chrome is recommended'),
        ]
    )

    module_chapter(doc,12,'PKD','PERFORMANCE KPI DASHBOARD',
        purpose_paras=[
            'The Performance KPI Dashboard tracks the municipality\'s actual performance against SDBIP (Service Delivery and Budget Implementation Plan) targets on a quarterly basis. It provides a structured, evidence-based mechanism for directors to submit performance actuals and for the MM and Council to monitor IDP implementation in real time.',
            'Performance management failures — including KPIs that are not SMART, targets not captured, or actuals not submitted — are consistently raised by AGSA. The PKD module creates a permanent, auditable performance record that supports the annual Section 46 performance report.',
        ],
        legislation='MSA Section 41 (performance management systems), Municipal Planning and Performance Management Regulations 2001, MFMA Section 72 (mid-year performance assessment), SDBIP requirements under MSA.',
        users=[
            ('Performance Manager','Captures all SDBIP KPIs; manages annual target-setting; coordinates quarterly submissions; generates reports'),
            ('Directors','Submit quarterly actual performance against their directorate\'s KPIs; capture corrective action plans for Red KPIs'),
            ('Municipal Manager','Reviews consolidated performance; escalates persistently red KPIs; approves Section 46 report'),
            ('Council / MPAC','Quarterly performance oversight; approves the annual Section 46 performance report'),
            ('Mayor','Service delivery overview; IDP implementation progress; political accountability for red KPIs'),
        ],
        features=[
            ('KPI Register','All SDBIP-linked KPIs with annual targets and quarterly milestone targets per KPI'),
            ('Quarterly Capture','Directors submit actual performance each quarter through a structured input form'),
            ('RAG Status','Red = below 50% of quarterly target; Amber = 50-79%; Green = 80% and above — calculated automatically'),
            ('Directorate Scorecards','Performance summary for each directorate showing overall achievement rate and number of Red/Amber/Green KPIs'),
            ('Corrective Action Plans','Mandatory for all Red KPIs — capture what will be done, by whom, and by when'),
            ('Section 46 Report Feed','Annual performance data auto-compiled from all four quarters for the statutory Section 46 report'),
            ('Year-on-Year Comparison','Current year performance compared to the same KPI in the prior financial year'),
        ],
        steps=[
            (1,'SETUP (Financial Year Start) — Performance Manager captures all SDBIP KPIs including annual targets and quarterly milestone targets'),
            (2,'QUARTER-END — Each directorate Director logs in > Click "Performance KPIs (PKD)" > Select their directorate'),
            (3,'SUBMIT ACTUALS — Director enters the actual performance figure for each KPI for the quarter just ended'),
            (4,'REVIEW STATUS — System automatically calculates the RAG status based on percentage of quarterly target achieved'),
            (5,'CORRECTIVE ACTION — For every Red KPI: Director captures a corrective action plan (what action, responsible person, completion date)'),
            (6,'COMPLETENESS CHECK — Performance Manager reviews all directorate submissions for completeness; flags any missing submissions'),
            (7,'MM REVIEW — MM reviews the consolidated performance report; escalates persistent Red KPIs to the responsible Director'),
            (8,'QUARTERLY REPORT — Performance Manager generates the quarterly performance report for Council/MPAC submission'),
            (9,'YEAR-END — At financial year end, the system auto-compiles the Section 46 annual performance report from all four quarterly submissions'),
        ],
        tip_label='SMART KPIs Required:',
        tip_text='All KPIs must be SMART: Specific, Measurable, Achievable, Relevant, and Time-bound. Vague KPIs such as "improve service delivery" or "enhance governance" cannot be objectively assessed by AGSA and will be raised as a finding. Every Red KPI must have a corrective action plan in the system — "under review" or "to be determined" are not acceptable and will be questioned by MPAC. The Section 46 annual performance report must be submitted to Council by 31 January following the financial year end.',
        errors=[
            ('"KPI showing Red despite good actual performance"','Annual target or unit of measure was set incorrectly','Review and correct the KPI target with Performance Manager approval; document the correction in the audit trail'),
            ('"Director cannot see their KPIs"','KPIs not assigned to the correct directorate during setup','Performance Manager must edit each affected KPI and assign it to the correct directorate'),
            ('"Section 46 report missing KPIs"','Not all KPIs have been flagged as SDBIP-linked during capture','Edit each missing KPI and mark it as SDBIP-linked; only SDBIP-linked KPIs are included in the annual report'),
        ]
    )

    module_chapter(doc,13,'ERP','EXECUTIVE REPORTING PORTAL',
        purpose_paras=[
            'The Executive Reporting Portal provides the Mayor, Municipal Manager, CFO, and Council committees with a single consolidated governance and financial intelligence view. It draws live data from all 12 operational modules to produce committee-ready reports, AI-generated narrative assessments, and trend analyses — eliminating the need to compile governance reports manually from multiple, disconnected sources.',
            'The ERP is designed for use before Council meetings, MPAC sittings, and Executive Committee briefings. It transforms module-level data into decision-ready intelligence, enabling senior officials to enter governance discussions fully informed.',
        ],
        legislation='MFMA Section 129 (municipal annual report), MSA Section 46 (annual performance report), MFMA Section 71/72 (financial reports), COGTA governance reporting requirements for municipalities.',
        users=[
            ('Mayor','Governance overview before Council meetings; service delivery snapshot; political reporting preparation'),
            ('Municipal Manager','Comprehensive governance report for management committees and Executive Committee'),
            ('CFO','Financial governance overview; trend analysis for the Finance and Budget Committee'),
            ('Council','Oversight of overall municipal governance performance; accountability to residents'),
            ('MPAC','Municipal Public Accounts Committee — uses ERP data for forensic-level accountability questioning'),
        ],
        features=[
            ('Consolidated Dashboard','Key governance metrics from all 12 modules in a single, integrated view'),
            ('AI-Generated Narrative','Written assessment identifying the top governance risks, trends, and recommended management actions'),
            ('Period Selection','Select any financial year quarter for current or historical governance reporting'),
            ('Committee Report Pack','One-click generation of a formatted, committee-ready governance report covering all modules'),
            ('Trend Analysis','12-month trend charts for open findings, IFWE total, SCM compliance rate, and KPI achievement'),
            ('Comparative Benchmarking','Your municipality\'s key metrics compared to peer group averages where benchmark data is available'),
            ('PDF and Word Export','Export report packs in both formats for Council agenda inclusion and MPAC distribution'),
        ],
        steps=[
            (1,'Log in as Mayor, MM, or CFO > Click "Executive Portal (ERP)" in the sidebar'),
            (2,'Select the reporting period from the period dropdown (quarter and financial year)'),
            (3,'Review the consolidated governance dashboard — all key metrics from every module in a single view'),
            (4,'Read the AI-generated narrative — it summarises the governance position and highlights the top 3 risks in plain language'),
            (5,'Edit or supplement the narrative if needed before presenting to Council — personalise it for the specific meeting context'),
            (6,'Click "Generate Report Pack" — the system compiles a full, formatted committee-ready governance document'),
            (7,'Export the report pack as PDF > add to the Council or MPAC agenda pack > distribute to members at least 48 hours before the meeting'),
            (8,'Archive the report in the system — it remains available for AGSA inspection as evidence of ongoing governance oversight'),
        ],
        tip_label='Meeting Preparation:',
        tip_text='Use the ERP portal at least two days before every Council meeting, MPAC sitting, or Executive Committee — this provides time to investigate any unexpected metric before having to explain it in a formal Council sitting. The AI narrative is a starting point, not a final product; always review, personalise, and verify it before presenting. MPAC members will ask detailed, specific questions — ensure you can drill down from ERP summary metrics into the underlying module records to support every number.',
        errors=[
            ('"Report shows no data for the selected period"','Underlying modules have no data captured for the selected period','Ensure all 12 modules have current data entered for the selected quarter before generating the ERP report'),
            ('"AI narrative not generating"','Connection timeout or insufficient data across modules','Refresh the page; ensure at least 3 modules have data for the selected period and try again'),
            ('"PDF export shows empty pages"','Chart rendering issue in the browser','Use Google Chrome for the best PDF rendering results; avoid Internet Explorer'),
        ]
    )

    module_chapter(doc,14,'SET','ENTITY SETTINGS',
        purpose_paras=[
            'The Entity Settings module is the configuration and personalisation centre for your PSFGS deployment. It stores all municipality-specific information — from the official entity name and logo to leadership names and contact details — and makes this information available throughout the entire system, including on the login page, in all auto-generated reports, and in the sidebar.',
            'Proper configuration of Entity Settings ensures that every report, letter, and document generated by PSFGS carries the correct municipal identity, leadership information, and official branding.',
        ],
        legislation='Not applicable — this is an administrative configuration module. Access is restricted to System_Admin only.',
        users=[
            ('System_Admin','The only role with access to Entity Settings. Typically the municipal IT Manager or designated PSFGS Administrator.'),
        ],
        features=[
            ('Entity Name and Abbreviation','Full official municipality name (e.g. Buffalo City Metropolitan Municipality) and abbreviation (e.g. BCMM)'),
            ('Logo Upload','Upload the official municipal crest or logo in PNG or JPG format (maximum 2MB) — appears in the sidebar and login page immediately after saving'),
            ('Demarcation Code','Official Statistics South Africa demarcation code (e.g. BUF for Buffalo City Metropolitan Municipality)'),
            ('Province and Category','Select the province and municipality category (Metropolitan / District / Local)'),
            ('Leadership Names','Municipal Manager, CFO, Mayor, and Speaker names — used automatically in generated reports and delegation letters'),
            ('Contact Details','Phone, fax, email, website, physical address, and postal address'),
            ('Financial Year End','Set the municipality\'s financial year end month (default: June for all MFMA municipalities)'),
        ],
        steps=[
            (1,'Log in as System_Admin > Click "Settings" at the bottom of the left sidebar'),
            (2,'GENERAL TAB — Enter Entity Name, Abbreviation, Demarcation Code, Province, and Municipality Category'),
            (3,'LOGO UPLOAD — Click "Upload Logo" > Select a PNG or JPG file (max 2MB) > Preview displays immediately > Logo will appear in sidebar and login page after saving'),
            (4,'CONTACT TAB — Enter Phone, Fax, Email, Website, Physical Address, and Postal Address'),
            (5,'LEADERSHIP TAB — Enter the full names of the Municipal Manager, CFO, Mayor, and Speaker'),
            (6,'Click "Save Settings" — all information is immediately active and visible across the entire system'),
            (7,'Verify that the sidebar now displays the new logo and updated entity name'),
            (8,'Verify that the login page now shows the new logo and entity name for a clean first-impression to all users'),
        ],
        tip_label='Maintenance Required:',
        tip_text='Update all leadership names at the start of each financial year and immediately whenever a new appointment is made — these names appear in auto-generated Section 71 reports, consequence management letters, delegation documents, and Council report packs. Outdated names in official documents are unprofessional and create audit trail inconsistencies. Only System_Admin can access Settings — this account must be protected with a strong password and must not be shared.',
        errors=[
            ('"Logo not appearing after upload"','File size exceeds the 2MB limit','Compress the image to below 2MB using a free online image compressor and re-upload'),
            ('"Logo shows broken image icon in sidebar"','Unsupported file format was uploaded','Use only PNG or JPG file formats; avoid TIFF, BMP, or SVG formats'),
            ('"Leadership names not appearing in reports"','Name fields were left blank or Save was not clicked','Complete all leadership name fields and click Save Settings; verify by opening a generated report'),
        ]
    )

    # ── APPENDIX A: PERMISSIONS MATRIX ────────────────────────────────────────
    doc.add_page_break()
    chapter_banner(doc,'APPENDIX A: USER ROLES AND PERMISSIONS MATRIX')
    para(doc,'The table below shows the access level for each of the eight PSFGS user roles across all 13 modules. Full = Full read/write access. Write = Write access within permitted scope. Read = Read-only access. None = No access. Admin = Full administrative access.',bef=0,aft=80)

    mods=['DASH','AFT','AVS','BVM','IFW','SCC','CMT','DAR','PCR','S71','PKD','ERP','SET']
    roles_matrix=[
        ('System_Admin',   ['Admin']*13),
        ('Acctg_Officer',  ['Full','Write','Read','Write','Write','Read','Full','Full','Read','Read','Read','Full','None']),
        ('CFO',            ['Full','Write','Read','Full','Full','Read','Read','Read','Read','Full','Read','Full','None']),
        ('Director',       ['Read','Write','Write','Read','Read','Read','Read','Read','Read','None','Write','Read','None']),
        ('Manager',        ['Read','Write','Write','Read','None','Write','None','None','Read','None','Write','None','None']),
        ('Officer',        ['Read','Write','Write','None','None','None','None','None','None','None','None','None','None']),
        ('Auditor',        ['Read','Full','Read','Read','Read','Read','Read','Read','Full','Read','Read','Read','None']),
        ('Councillor',     ['Read','None','None','Read','Read','None','Read','Read','Read','Read','Read','Read','None']),
    ]
    col_widths=[Cm(3.2)]+[Cm(1.02)]*13
    at=doc.add_table(len(roles_matrix)+1,14)
    at.alignment=WD_TABLE_ALIGNMENT.CENTER
    # header
    for i,(h,w) in enumerate(zip(['Role']+mods,col_widths)):
        c=at.rows[0].cells[i]; c.width=w; shade(c,DB); no_border(c); cell_mar(c,40,40,40,40)
        p=c.add_paragraph(); sp(p,0,0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,h,True,7.5,WH)
    access_colours={'Admin':DB,'Full':DG,'Write':MB,'Read':GR,'None':RGBColor(0xCC,0xCC,0xCC)}
    access_bg={'Admin':LB,'Full':DG2,'Write':LB,'Read':LG,'None':WH}
    for ri,(role,perms) in enumerate(roles_matrix):
        for ci,(v,w) in enumerate(zip([role]+perms,col_widths)):
            c=at.rows[ri+1].cells[ci]; c.width=w
            bg=LG if ri%2==0 else WH
            if ci>0: bg=access_bg.get(v,WH)
            shade(c,bg); no_border(c); cell_mar(c,40,40,40,40)
            p=c.add_paragraph(); sp(p,0,0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            col_=access_colours.get(v,BD) if ci>0 else DB
            run(p,v,ci==0,7.5,col_)
    doc.add_paragraph()

    # ── APPENDIX B: GLOSSARY ──────────────────────────────────────────────────
    doc.add_page_break()
    chapter_banner(doc,'APPENDIX B: GLOSSARY OF TERMS')
    glossary=[
        ('AGSA','Auditor-General of South Africa — the Supreme Audit Institution constitutionally mandated to audit all national, provincial, and local government entities'),
        ('AFS','Annual Financial Statements — the statutory financial statements that every municipality must prepare and submit within two months of financial year end'),
        ('BEE','Black Economic Empowerment — preferential procurement requirements applied under the Preferential Procurement Policy Framework Act'),
        ('CFO','Chief Financial Officer — the senior official responsible for sound financial management under MFMA Section 80'),
        ('COGTA','Department of Cooperative Governance and Traditional Affairs — the national department responsible for overseeing and supporting municipalities'),
        ('CSD','Central Supplier Database — the National Treasury\'s centralised database of all registered government suppliers; all suppliers must be registered on CSD'),
        ('DAR','Delegation of Authority Register — the formal record of all powers delegated by Council to the Municipal Manager and all sub-delegations'),
        ('GRAP','Generally Recognised Accounting Practice — the accounting standards applicable to all South African public sector entities, analogous to IFRS for government'),
        ('IDP','Integrated Development Plan — the 5-year strategic plan that every municipality must prepare, reviewing it annually and adopting a new IDP every term'),
        ('IFWE','Irregular, Fruitless and Wasteful Expenditure — three categories of expenditure defined in MFMA Sections 32, 33, and 34 that must be identified, investigated, and disclosed'),
        ('MFMA','Municipal Finance Management Act 56 of 2003 — the primary legislation governing all aspects of municipal financial management in South Africa'),
        ('MPAC','Municipal Public Accounts Committee — the Council committee established to scrutinise the annual report and hold the administration accountable'),
        ('MSA','Municipal Systems Act 32 of 2000 — legislation governing municipal systems, integrated planning, performance management, and service delivery'),
        ('NT','National Treasury — the national department responsible for fiscal policy, intergovernmental financial relations, and oversight of MFMA compliance'),
        ('PFMA','Public Finance Management Act 1 of 1999 — the financial management legislation governing national and provincial government departments (not municipalities)'),
        ('PPPFA','Preferential Procurement Policy Framework Act 5 of 2000 — legislation requiring government entities to apply preferential procurement to advance BEE objectives'),
        ('PSFGS','Public Sector Financial Governance Suite — the integrated software platform described in this training manual'),
        ('RAG','Red-Amber-Green — the colour-coded performance status system used in the PKD module (Red below 50%, Amber 50-79%, Green 80%+)'),
        ('SDBIP','Service Delivery and Budget Implementation Plan — the annual operational plan linking the IDP and budget to measurable quarterly KPI targets'),
        ('SCM','Supply Chain Management — the procurement, logistics, and disposal management system that every municipality must maintain under the MFMA SCM Regulations'),
        ('S71','Section 71 — refers to MFMA Section 71 which requires the CFO to submit monthly financial statements to National Treasury by the 10th of each month'),
        ('SAPS','South African Police Service — must be notified within 24 hours of any theft or suspected theft of municipal assets or funds'),
    ]
    gt=doc.add_table(len(glossary)+1,2)
    gt.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(['Term','Definition'],[Cm(3.5),Cm(15)])):
        c=gt.rows[0].cells[i]; c.width=w; shade(c,DB); no_border(c); cell_mar(c,60,60,80,80)
        p=c.add_paragraph(); sp(p,0,0); run(p,h,True,8.5,WH)
    for ri,(term,defn) in enumerate(glossary):
        for ci,(v,w) in enumerate(zip([term,defn],[Cm(3.5),Cm(15)])):
            c=gt.rows[ri+1].cells[ci]; c.width=w
            shade(c,LB if ri%2==0 else WH); no_border(c); cell_mar(c,50,50,80,80)
            p=c.add_paragraph(); sp(p,0,0)
            run(p,v,ci==0,8.5,DB if ci==0 else BD)
    doc.add_paragraph()

    # ── APPENDIX C: SUPPORT ───────────────────────────────────────────────────
    doc.add_page_break()
    chapter_banner(doc,'APPENDIX C: SUPPORT AND CONTACTS')

    section_label(doc,'Getting Technical Support')
    para(doc,'For any technical issues with PSFGS, contact your municipality\'s designated System Administrator as the first point of contact. The System Administrator has access to all configuration settings, user management functions, and system logs, and should be able to resolve most login, access, and configuration issues.',bef=0)
    para(doc,'If the issue cannot be resolved by your System Administrator, contact Glance Management Technologies directly using the details below.',bef=0)

    section_label(doc,'Glance Management Technologies (Pty) Ltd — Product Support')
    support=[
        ('Email','support@glancemanagement.co.za'),
        ('Website','www.psfgs.co.za'),
        ('Business Hours','Monday to Friday, 08:00 to 17:00 (South Africa Standard Time)'),
        ('Response Time','Critical issues: within 4 hours. Standard queries: within 1 business day'),
    ]
    st=doc.add_table(len(support),2)
    st.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,(label,val) in enumerate(support):
        for ci,(v,w) in enumerate(zip([label,val],[Cm(5),Cm(12.5)])):
            c=st.rows[ri].cells[ci]; c.width=w
            shade(c,LB if ri%2==0 else WH); no_border(c); cell_mar(c,50,50,80,80)
            p=c.add_paragraph(); sp(p,0,0); run(p,v,ci==0,9,DB if ci==0 else BD)
    doc.add_paragraph()

    section_label(doc,'MFMA Compliance Reference Resources')
    for item in [
        ('National Treasury — Local Government Budget Analysis','www.treasury.gov.za/legislation/mfma'),
        ('COGTA — Municipal Support and Interventions','www.cogta.gov.za'),
        ('AGSA — Audit Reports and Management Letters','www.agsa.co.za'),
        ('SALGA — Municipal Governance Support','www.salga.org.za'),
    ]: bullet(doc,item[1],bold_prefix=item[0])

    section_label(doc,'Training and User Onboarding')
    para(doc,'For onboarding assistance, system demonstrations, or hands-on training for your municipality\'s PSFGS team, contact Glance Management Technologies to arrange a training session. Training is available as:')
    for item in [
        'On-site full-day training workshop (recommended for initial deployment)',
        'Remote video training sessions (per module or full suite)',
        'Self-paced online training materials (available on www.psfgs.co.za)',
        'Super-user training for System Administrators and Power Users',
    ]: bullet(doc,item)

    para(doc,'',bef=60,aft=0)
    para(doc,'This training manual is version 1.0. For the latest version and any updates, visit www.psfgs.co.za/documentation',False,8.5,GR)
    para(doc,'Copyright 2026 Glance Management Technologies (Pty) Ltd. All rights reserved. This document is for official use by licensed PSFGS municipalities only.',False,8,GR,bef=0)

    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__=='__main__':
    build()
